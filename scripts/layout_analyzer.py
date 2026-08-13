"""Read reusable layout metrics from a reference Word document.

The analyzer deliberately learns measurable page and paragraph formatting;
it never copies document text into the preset. Complex drawings, logos,
headers, and floating objects remain outside the automatic-learning scope.
"""

from __future__ import annotations

from collections import Counter
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


class LayoutAnalysisError(RuntimeError):
    """Raised when a reference file cannot provide a safe reusable layout."""


@dataclass(frozen=True)
class LayoutAnalysis:
    settings: dict
    summary: str
    warnings: tuple[str, ...]


_ALIGNMENTS = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    WD_ALIGN_PARAGRAPH.DISTRIBUTE: "justify",
}


def _points(value, fallback=0.0):
    if value is None:
        return float(fallback)
    if hasattr(value, "pt"):
        return float(value.pt)
    try:
        return float(value)
    except (TypeError, ValueError):
        return float(fallback)


def _length_cm(value, fallback):
    try:
        return round(float(value.cm), 2)
    except (TypeError, ValueError, AttributeError):
        return float(fallback)


def _length_mm(value, fallback):
    try:
        return round(float(value.mm), 2)
    except (TypeError, ValueError, AttributeError):
        return float(fallback)


def _paragraph_value(paragraph, attribute, default=None):
    value = getattr(paragraph.paragraph_format, attribute, None)
    if value is not None:
        return value
    style = getattr(paragraph, "style", None)
    style_format = getattr(style, "paragraph_format", None)
    value = getattr(style_format, attribute, None) if style_format is not None else None
    return default if value is None else value


def _run_font_names(run, paragraph):
    east_asia = ""
    ascii_name = ""
    r_pr = run._r.rPr
    if r_pr is not None:
        fonts = r_pr.find(qn("w:rFonts"))
        if fonts is not None:
            east_asia = fonts.get(qn("w:eastAsia")) or ""
            ascii_name = fonts.get(qn("w:ascii")) or fonts.get(qn("w:hAnsi")) or ""
    style_font = getattr(getattr(paragraph, "style", None), "font", None)
    direct_name = run.font.name or ""
    style_name = getattr(style_font, "name", None) or ""
    font_cn = east_asia or direct_name or style_name or "宋体"
    font_en = ascii_name or direct_name or style_name or "Times New Roman"
    return font_cn, font_en


def _style_ppr(paragraph):
    style = getattr(paragraph, "style", None)
    element = getattr(style, "element", None)
    return getattr(element, "pPr", None)


def _indent_points(paragraph, size):
    """First-line indent in points.

    python-docx only reads ``w:firstLine`` (twips); Chinese Word/WPS documents
    commonly store "首行缩进 N 字符" as ``w:firstLineChars`` (1/100 char), which
    python-docx reports as ``None``.  Read the raw attributes first, paragraph
    level then style level, and fall back to the python-docx chain last.
    """
    for element in (paragraph._p.pPr, _style_ppr(paragraph)):
        if element is None:
            continue
        ind = element.find(qn("w:ind"))
        if ind is None:
            continue
        chars = ind.get(qn("w:firstLineChars"))
        if chars:
            try:
                return max(0.0, float(chars) / 100.0) * max(size, 1.0)
            except (TypeError, ValueError):
                pass
        if ind.get(qn("w:hangingChars")) or ind.get(qn("w:hanging")):
            return 0.0
        first_line = ind.get(qn("w:firstLine"))
        if first_line:
            try:
                # w:firstLine is stored in twentieths of a point.
                return max(0.0, float(first_line) / 20.0)
            except (TypeError, ValueError):
                pass
    return max(0.0, _points(_paragraph_value(paragraph, "first_line_indent"), 0.0))


def _leading_space_chars(text):
    """Approximate width of leading whitespace, in full-width characters."""
    count = 0.0
    for char in str(text or ""):
        if char == "　":
            count += 1.0
        elif char == " ":
            count += 0.5
        elif char == "\t":
            count += 2.0
        else:
            break
    return count


def _snap_indent(indent, size, text=""):
    """Snap a measured indent to whole-character conventions.

    Hand-formatted documents mix real indents with full-width spaces; snapping
    keeps tiny measurement noise (e.g. 0.6 字残留) and space-based indents from
    becoming the learned value.
    """
    size = max(float(size or 0), 1.0)
    ratio = float(indent or 0) / size
    if 1.5 <= ratio <= 2.6:
        return round(2 * size, 1)
    if 0.9 <= ratio < 1.5:
        return round(size, 1)
    if 0.5 <= ratio < 0.9:
        # 0.5–0.9 字多半是列表/制表位残留，不是真实缩进
        return 0.0
    if ratio < 0.5 and _leading_space_chars(text) >= 1.5:
        return round(2 * size, 1)
    return round(max(float(indent or 0), 0.0), 1)


def _paragraph_format(paragraph, *, fallback_size=12.0):
    visible_runs = [run for run in paragraph.runs if run.text.strip()]
    # Use the longest visible run as the paragraph's representative text style.
    # Regulation paragraphs often bold only the short ``第一条`` marker; using
    # the first run would incorrectly learn that the whole paragraph is bold.
    run = (
        max(visible_runs, key=lambda item: len(item.text.strip()))
        if visible_runs else (paragraph.runs[0] if paragraph.runs else None)
    )
    style_font = getattr(getattr(paragraph, "style", None), "font", None)
    if run is not None:
        font_cn, font_en = _run_font_names(run, paragraph)
        size = _points(run.font.size, _points(getattr(style_font, "size", None), fallback_size))
        bold_value = run.font.bold
    else:
        font_cn = getattr(style_font, "name", None) or "宋体"
        font_en = getattr(style_font, "name", None) or "Times New Roman"
        size = _points(getattr(style_font, "size", None), fallback_size)
        bold_value = None
    if bold_value is None:
        bold_value = getattr(style_font, "bold", None)

    alignment = _paragraph_value(paragraph, "alignment")
    line_spacing_value = _paragraph_value(paragraph, "line_spacing")
    if hasattr(line_spacing_value, "pt"):
        line_spacing = round(float(line_spacing_value.pt), 1)
    elif isinstance(line_spacing_value, (int, float)):
        line_spacing = round(float(line_spacing_value) * max(size, 1.0), 1)
    else:
        line_spacing = round(max(size * 1.5, size), 1)

    return {
        "font_cn": font_cn,
        "font_en": font_en,
        "size": round(max(size, 5.0), 1),
        "bold": bool(bold_value),
        "align": _ALIGNMENTS.get(alignment, "left"),
        "indent": _snap_indent(_indent_points(paragraph, size), size, paragraph.text),
        "line_spacing": line_spacing,
        "space_before": round(_points(_paragraph_value(paragraph, "space_before"), 0.0), 1),
        "space_after": round(_points(_paragraph_value(paragraph, "space_after"), 0.0), 1),
    }


def _explicit_style_role(paragraph):
    style = getattr(paragraph, "style", None)
    style_name = str(getattr(style, "name", "") or "").lower().replace(" ", "")
    style_id = str(getattr(style, "style_id", "") or "").lower().replace(" ", "")
    combined = f"{style_name}|{style_id}"
    if any(value in combined for value in ("heading1", "标题1", "标题一")):
        return "heading1"
    if any(value in combined for value in ("heading2", "标题2", "标题二")):
        return "heading2"
    if any(value in combined for value in ("heading3", "标题3", "标题三")):
        return "heading3"
    if any(value in combined for value in ("heading4", "标题4", "标题四")):
        return "heading4"
    if style_id == "title" or style_name in {"title", "标题", "文档标题"}:
        return "title"
    return ""


def _text_role(text):
    text = str(text or "").strip()
    if re.match(r"^[一二三四五六七八九十]+、", text):
        return "heading1"
    if re.match(r"^[（(][一二三四五六七八九十]+[）)]", text):
        return "heading2"
    # 与 formatter._looks_like_heading3 保持一致：只认 "1."，排除 "1.1" 等版本号
    if re.match(r"^\d+\.\s*\S", text) and not re.match(r"^\d+\.\d", text) and len(text) < 60:
        return "heading3"
    if re.match(r"^[（(]\d+[）)]", text) and len(text) < 70:
        return "heading4"
    return ""


def _choose_title(paragraphs, body_size):
    explicit = [item for item in paragraphs if item["role"] == "title"]
    if explicit:
        return explicit[0]
    # 常识约束：标题在开头、较短、不会比正文小，且总有点「标题相」（居中或加粗）。
    # 不满足就宁可不学——否则普通正文会被学成标题。
    candidates = [
        item for item in paragraphs[:10]
        if len(item["text"]) <= 40
        and item["format"]["size"] >= body_size - 0.5
        and (item["format"]["align"] == "center" or item["format"]["bold"])
    ]
    if not candidates:
        return None
    centered = [item for item in candidates if item["format"]["align"] == "center"]
    pool = centered or candidates
    return max(
        pool,
        key=lambda item: (
            item["format"]["size"],
            bool(item["format"]["bold"]),
            -item["index"],
        ),
    )


def _format_signature(value):
    return tuple((key, value.get(key)) for key in (
        "font_cn", "font_en", "size", "bold", "align", "indent",
        "line_spacing", "space_before", "space_after",
    ))


def _choose_body(paragraphs):
    pool = [
        item for item in paragraphs
        if not item["role"] and len(item["text"]) >= 8
    ]
    if not pool:
        pool = [item for item in paragraphs if not item["role"]]
    if not pool:
        return None
    # 先按「中文字体 + 字号」聚类，避免手写文档里签名碎片化导致多数派失真；
    # 再在胜出簇内取完整签名的代表格式。
    clusters = Counter(
        (item["format"].get("font_cn"), item["format"].get("size")) for item in pool
    )
    top_cluster = clusters.most_common(1)[0][0]
    members = [
        item for item in pool
        if (item["format"].get("font_cn"), item["format"].get("size")) == top_cluster
    ]
    signature = Counter(_format_signature(item["format"]) for item in members).most_common(1)[0][0]
    return next(item for item in members if _format_signature(item["format"]) == signature)


def _footer_has_page_number(doc):
    for section in doc.sections:
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            xml = footer._element.xml
            if re.search(r"\bPAGE\b", xml, flags=re.IGNORECASE):
                return True
    return False


def collect_reference_paragraph_samples(path, *, limit=120):
    """Read text and measurable formatting for an explicit local analysis step.

    Unlike ``analyze_reference_layout``, this result contains paragraph text and
    must never be persisted into a preset. Callers should show a privacy notice
    before sending any sample to an external model.
    """
    path = Path(path).expanduser()
    if not path.is_file() or path.suffix.lower() != ".docx":
        raise LayoutAnalysisError("请选择 .docx 格式的参考 Word 文档。")
    try:
        doc = Document(path)
    except Exception as exc:
        raise LayoutAnalysisError(f"无法读取参考 Word：{exc}") from exc
    samples = []
    max_items = max(1, min(int(limit or 120), 300))
    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        if not text:
            continue
        style = getattr(paragraph, "style", None)
        samples.append({
            "index": index,
            "text": text,
            "style_name": str(getattr(style, "name", "") or ""),
            "format": _paragraph_format(paragraph),
        })
        if len(samples) >= max_items:
            break
    if not samples:
        raise LayoutAnalysisError("参考 Word 没有可分析的正文段落。")
    return samples


def analyze_reference_layout(path, *, base_settings=None):
    """Return a generic-layout preset learned from a `.docx` reference."""
    path = Path(path).expanduser()
    if not path.is_file() or path.suffix.lower() != ".docx":
        raise LayoutAnalysisError("请选择 .docx 格式的参考 Word 文档。")
    try:
        doc = Document(path)
    except Exception as exc:
        raise LayoutAnalysisError(f"无法读取参考 Word：{exc}") from exc
    if not doc.sections:
        raise LayoutAnalysisError("参考 Word 中没有可读取的页面设置。")

    settings = deepcopy(base_settings or {})
    section = doc.sections[0]
    width_mm = _length_mm(section.page_width, 210.0)
    height_mm = _length_mm(section.page_height, 297.0)
    orientation = "landscape" if width_mm > height_mm else "portrait"
    settings["layout_mode"] = "generic"
    settings["page"] = {
        **settings.get("page", {}),
        "top": _length_cm(section.top_margin, 2.54),
        "bottom": _length_cm(section.bottom_margin, 2.54),
        "left": _length_cm(section.left_margin, 2.54),
        "right": _length_cm(section.right_margin, 2.54),
        "width_mm": round(min(width_mm, height_mm), 2),
        "height_mm": round(max(width_mm, height_mm), 2),
        "orientation": orientation,
    }

    records = []
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        explicit_role = _explicit_style_role(paragraph)
        records.append({
            "index": index,
            "text": text,
            "role": explicit_role or _text_role(text),
            "format": _paragraph_format(paragraph),
        })
    if not records:
        raise LayoutAnalysisError("参考 Word 没有可分析的正文段落。")

    body = _choose_body(records)
    if body:
        body_size = body["format"].get("size", 12)
    else:
        body_size = _points(settings.get("body", {}).get("size"), 12)
    title = _choose_title(records, body_size)
    if title:
        settings["title"] = title["format"]
    if body:
        settings["body"] = body["format"]
    for role in ("heading1", "heading2", "heading3", "heading4"):
        sample = next((item for item in records if item["role"] == role), None)
        if sample:
            settings[role] = sample["format"]

    body_format = deepcopy(settings.get("body", {}))
    for role in ("recipient", "signature", "date", "attachment", "closing"):
        settings[role] = deepcopy(body_format)
    settings["table"] = {
        **settings.get("table", {}),
        "optimize": False,
        "before_table_blank_line": False,
        "after_table_blank_line": False,
        "header_bold": False,
        "smart_align": False,
    }
    table_paragraphs = [
        paragraph
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
        if paragraph.text.strip()
    ]
    if table_paragraphs:
        table_format = _paragraph_format(table_paragraphs[0], fallback_size=10.5)
        settings["table"].update({
            "font_cn": table_format["font_cn"],
            "font_en": table_format["font_en"],
            "size": table_format["size"],
            "bold": False,
            "line_spacing": table_format["line_spacing"],
            "header_bold": table_format["bold"],
        })

    settings["space_handling"] = "keep_all"
    settings["deep_clean"] = False
    settings["remove_background"] = False
    settings["split_heading_at_punct"] = False
    settings["first_line_bold"] = False
    settings["bold_serial"] = False
    settings["page_number"] = _footer_has_page_number(doc)
    settings["layout_source"] = {
        "file_name": path.name,
        "paragraph_count": len(records),
        "table_count": len(doc.tables),
        "section_count": len(doc.sections),
    }

    learned = ["页面", "主标题", "正文"]
    learned.extend(role for role in ("heading1", "heading2", "heading3", "heading4") if any(item["role"] == role for item in records))
    warnings = []
    if len(records) < 12 and len(doc.tables) >= 1:
        warnings.append("该文档大部分内容可能在表格中；自动学习对表格类文档可能不准，表格部分请排版后检查。")
    if len(doc.sections) > 1:
        warnings.append("参考文档含多个分节；当前模板读取第一节页面参数，处理时仍保留输入文档原有分节。")
    if any(
        any(paragraph.text.strip() for paragraph in section.header.paragraphs)
        or "<w:drawing" in section.header._element.xml
        or "<w:pict" in section.header._element.xml
        for section in doc.sections
    ):
        warnings.append("页眉、徽标、文本框和浮动图形不会复制，只读取可复用的文字与页面指标。")
    summary = f"已读取{len(records)}个正文段落、{len(doc.tables)}个表格；提取：{'、'.join(learned)}。"
    return LayoutAnalysis(settings=settings, summary=summary, warnings=tuple(warnings))
