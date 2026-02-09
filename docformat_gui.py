#!/usr/bin/env python3
"""
公文格式处理工具 - 纸质感极简风格 v2
优化：更大图标、更好排版、卡片式选择
"""

import os
import sys
import threading
import tkinter as tk
from tkinter import filedialog, messagebox
from pathlib import Path

# 添加scripts目录到路径
SCRIPT_DIR = Path(__file__).parent / "scripts"
sys.path.insert(0, str(SCRIPT_DIR))

from scripts.analyzer import analyze_punctuation, analyze_numbering, analyze_paragraph_format, analyze_font
from scripts.punctuation import process_document as fix_punctuation
from scripts.formatter import format_document, PRESETS


# ===== 设计系统 =====
class Theme:
    # 纸质色调
    BG = '#FBF9F6'              # 温暖米白纸张
    CARD = '#FFFFFF'            # 纯白卡片
    CARD_ALT = '#F7F4EF'        # 米黄卡片（推荐区）
    INPUT_BG = '#F2EFE9'        # 输入框背景（稍深米色）
    
    # 陶土红
    PRIMARY = '#BC4B26'         # 朱砂/印泥色
    PRIMARY_HOVER = '#A3421F'   # 悬停加深
    PRIMARY_LIGHT = '#F9F0EC'   # 极淡红
    
    # 文字
    TEXT = '#2E2E2E'            # 深炭灰
    TEXT_SECONDARY = '#6B6B6B'  # 次要文字
    TEXT_MUTED = '#A0A0A0'      # 禁用/占位
    
    # 边框与分隔
    BORDER = '#E8E4DE'          # 温暖灰边框
    BORDER_LIGHT = '#F0EDE8'    # 更浅边框
    BORDER_SELECTED = '#BC4B26' # 选中边框
    
    # 日志区
    LOG_BG = '#1A1A1A'
    LOG_TEXT = '#C8C8C8'
    LOG_SUCCESS = '#7CB87C'
    LOG_WARNING = '#D4A656'
    LOG_ERROR = '#CF6B6B'
    
    # 字体 - 宋体优先
    FONT_SERIF = ('Noto Serif SC', 'Source Han Serif SC', 'SimSun', 'PMingLiU', 'serif')
    
    # 间距
    SPACE_XS = 4
    SPACE_SM = 8
    SPACE_MD = 16
    SPACE_LG = 24
    SPACE_XL = 40


def get_font(size=12, weight='normal'):
    """获取宋体字体"""
    return (Theme.FONT_SERIF[0], size, weight)


# ===== 配置管理 =====
import json

CONFIG_FILE = Path(__file__).parent / "custom_settings.json"

# 常用字体列表
COMMON_FONTS_CN = [
    '仿宋_GB2312', '仿宋', '宋体', '黑体', '楷体_GB2312', '楷体',
    '方正小标宋简体', '方正仿宋_GBK', '华文仿宋', '华文中宋'
]

COMMON_FONTS_EN = [
    'Times New Roman', 'Arial', 'Calibri', 'Cambria'
]

# 字号对照表
FONT_SIZES = [
    ('初号', 42), ('小初', 36), ('一号', 26), ('小一', 24),
    ('二号', 22), ('小二', 18), ('三号', 16), ('小三', 15),
    ('四号', 14), ('小四', 12), ('五号', 10.5), ('小五', 9),
]

DEFAULT_CUSTOM_SETTINGS = {
    'name': '自定义格式',
    'page': {'top': 3.46, 'bottom': 3.26, 'left': 2.8, 'right': 2.6},
    'title': {
        'font_cn': '方正小标宋简体', 'font_en': 'Times New Roman',
        'size': 22, 'bold': False, 'align': 'center', 'indent': 0,
        'line_spacing': 29.45, 'space_before': 0, 'space_after': 0
    },
    'recipient': {
        'font_cn': '仿宋_GB2312', 'font_en': 'Times New Roman',
        'size': 16, 'bold': False, 'align': 'left', 'indent': 0,
        'line_spacing': 29.45, 'space_before': 0, 'space_after': 0
    },
    'heading1': {
        'font_cn': '黑体', 'font_en': 'Times New Roman',
        'size': 16, 'bold': False, 'align': 'left', 'indent': 32,
        'line_spacing': 29.45, 'space_before': 0, 'space_after': 0
    },
    'heading2': {
        'font_cn': '楷体_GB2312', 'font_en': 'Times New Roman',
        'size': 16, 'bold': False, 'align': 'left', 'indent': 32,
        'line_spacing': 29.45, 'space_before': 0, 'space_after': 0
    },
    'heading3': {
        'font_cn': '仿宋_GB2312', 'font_en': 'Times New Roman',
        'size': 16, 'bold': False, 'align': 'left', 'indent': 32,
        'line_spacing': 29.45, 'space_before': 0, 'space_after': 0
    },
    'heading4': {
        'font_cn': '仿宋_GB2312', 'font_en': 'Times New Roman',
        'size': 16, 'bold': False, 'align': 'left', 'indent': 32,
        'line_spacing': 29.45, 'space_before': 0, 'space_after': 0
    },
    'body': {
        'font_cn': '仿宋_GB2312', 'font_en': 'Times New Roman',
        'size': 16, 'bold': False, 'align': 'justify',
        'indent': 32, 'line_spacing': 29.45, 'space_before': 0, 'space_after': 0
    },
    'signature': {
        'font_cn': '仿宋_GB2312', 'font_en': 'Times New Roman',
        'size': 16, 'bold': False, 'align': 'right', 'indent': 0,
        'line_spacing': 29.45, 'space_before': 0, 'space_after': 0
    },
    'date': {
        'font_cn': '仿宋_GB2312', 'font_en': 'Times New Roman',
        'size': 16, 'bold': False, 'align': 'right', 'indent': 0,
        'line_spacing': 29.45, 'space_before': 0, 'space_after': 0
    },
    'attachment': {
        'font_cn': '仿宋_GB2312', 'font_en': 'Times New Roman',
        'size': 16, 'bold': False, 'align': 'left', 'indent': 0,
        'line_spacing': 29.45, 'space_before': 0, 'space_after': 0
    },
    'closing': {
        'font_cn': '仿宋_GB2312', 'font_en': 'Times New Roman',
        'size': 16, 'bold': False, 'align': 'left', 'indent': 32,
        'line_spacing': 29.45, 'space_before': 0, 'space_after': 0
    },
    'table': {
        'font_cn': '仿宋_GB2312', 'font_en': 'Times New Roman',
        'size': 12, 'bold': False, 'line_spacing': 22,
        'first_line_indent': 0, 'header_bold': True
    },
    'first_line_bold': False,
    'page_number': True,
    'page_number_font': '宋体',
}


def load_custom_settings():
    """加载自定义设置"""
    if CONFIG_FILE.exists():
        try:
            with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
                return _merge_settings(DEFAULT_CUSTOM_SETTINGS, data)
        except Exception as e:
            print(f"[警告] 加载自定义设置失败: {e}，使用默认设置")
    return DEFAULT_CUSTOM_SETTINGS.copy()


def _merge_settings(defaults, custom):
    merged = {}
    for key, value in defaults.items():
        if key in custom:
            if isinstance(value, dict) and isinstance(custom.get(key), dict):
                merged[key] = _merge_settings(value, custom[key])
            else:
                merged[key] = custom[key]
        else:
            merged[key] = value
    return merged


def save_custom_settings(settings):
    """保存自定义设置"""
    try:
        with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
            json.dump(settings, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"[错误] 保存自定义设置失败: {e}")
        raise



# ===== 快速设置中，正文字体联动的元素 =====
BODY_FONT_GROUP = ['body', 'heading3', 'heading4', 'closing', 'attachment', 'signature', 'date']


class CustomSettingsDialog(tk.Toplevel):
    """自定义格式设置弹窗 - 快速设置 + 高级设置（可折叠）"""
    
    def __init__(self, parent, on_save=None):
        super().__init__(parent)
        
        self.on_save = on_save
        self.settings = load_custom_settings()
        self._adv_vars = {}  # 高级模式的变量存储
        
        # 窗口设置
        self.title("自定义格式设置")
        win_w, win_h = 1200, 860
        self.geometry(f"{win_w}x{win_h}")
        self.minsize(1040, 700)
        self.configure(bg=Theme.BG)
        self.resizable(True, True)
        
        # 模态窗口
        self.transient(parent)
        self.grab_set()
        self.protocol("WM_DELETE_WINDOW", self._on_close)
        
        # 居中显示
        self.update_idletasks()
        x = parent.winfo_x() + (parent.winfo_width() - win_w) // 2
        y = parent.winfo_y() + (parent.winfo_height() - win_h) // 2
        self.geometry(f"+{max(0, x)}+{max(0, y)}")
        
        self._create_widgets()
        self._load_values()
    
    # ==================== 界面构建 ====================
    
    def _create_widgets(self):
        """创建控件 - 快速设置 + 可折叠高级设置"""
        # ===== 顶部标题 + 按钮（固定） =====
        header = tk.Frame(self, bg=Theme.BG)
        header.pack(fill='x', padx=20, pady=(15, 5))
        
        tk.Label(
            header, text="⚙️ 自定义格式设置", font=get_font(16, 'bold'),
            bg=Theme.BG, fg=Theme.TEXT
        ).pack(side='left')
        
        # 保存按钮（顶部）
        save_top = tk.Frame(header, bg=Theme.PRIMARY, cursor='hand2')
        save_top.pack(side='right')
        save_top_label = tk.Label(
            save_top, text="  保存设置  ", font=get_font(12, 'bold'),
            bg=Theme.PRIMARY, fg='white', pady=6, cursor='hand2'
        )
        save_top_label.pack()
        for w in [save_top, save_top_label]:
            w.bind('<Button-1>', lambda e: self._save())
            w.bind('<Enter>', lambda e: (save_top.configure(bg=Theme.PRIMARY_HOVER), save_top_label.configure(bg=Theme.PRIMARY_HOVER)))
            w.bind('<Leave>', lambda e: (save_top.configure(bg=Theme.PRIMARY), save_top_label.configure(bg=Theme.PRIMARY)))
        
        cancel_top = tk.Label(
            header, text="取消", font=get_font(11),
            bg=Theme.BG, fg=Theme.TEXT_SECONDARY, cursor='hand2', padx=10
        )
        cancel_top.pack(side='right', padx=(0, 10))
        cancel_top.bind('<Button-1>', lambda e: self._on_close())
        
        # ===== 滚动区域 =====
        scroll_container = tk.Frame(self, bg=Theme.BG)
        scroll_container.pack(fill='both', expand=True, padx=5)
        
        self.canvas = tk.Canvas(scroll_container, bg=Theme.BG, highlightthickness=0)
        scrollbar = tk.Scrollbar(scroll_container, orient='vertical', command=self.canvas.yview)
        h_scrollbar = tk.Scrollbar(scroll_container, orient='horizontal', command=self.canvas.xview)
        
        self.canvas.configure(yscrollcommand=scrollbar.set, xscrollcommand=h_scrollbar.set)
        scrollbar.pack(side='right', fill='y')
        h_scrollbar.pack(side='bottom', fill='x')
        self.canvas.pack(side='left', fill='both', expand=True)
        
        self.content_frame = tk.Frame(self.canvas, bg=Theme.BG)
        self.canvas_window = self.canvas.create_window((0, 0), window=self.content_frame, anchor='nw')
        
        self.content_frame.bind('<Configure>', self._on_frame_configure)
        self.canvas.bind('<Configure>', self._on_canvas_configure)
        self.canvas.bind('<Enter>', lambda e: self._bind_mousewheel())
        self.canvas.bind('<Leave>', lambda e: self._unbind_mousewheel())
        
        main = self.content_frame
        pad_x = 15
        
        # ============================================================
        #  快速设置（始终显示）
        # ============================================================
        
        # --- 页面边距 ---
        self._create_section(main, "📄 页面边距 (cm)", pad_x)
        margin_frame = tk.Frame(main, bg=Theme.BG)
        margin_frame.pack(fill='x', pady=(0, 12), padx=pad_x)
        
        self.margin_vars = {}
        margins = [('top', '上'), ('bottom', '下'), ('left', '左'), ('right', '右')]
        for i, (key, label) in enumerate(margins):
            col = i % 4
            f = tk.Frame(margin_frame, bg=Theme.BG)
            f.grid(row=0, column=col, sticky='w', padx=(0, 15), pady=2)
            tk.Label(f, text=f"{label}:", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, anchor='e').pack(side='left')
            var = tk.StringVar(value=str(self.settings.get('page', {}).get(key, 2.5)))
            self.margin_vars[key] = var
            tk.Entry(f, textvariable=var, font=get_font(11), width=6, relief='solid', bd=1).pack(side='left', padx=3)
        
        # --- 标题格式 ---
        self._create_section(main, "📝 标题", pad_x)
        title_frame = tk.Frame(main, bg=Theme.BG)
        title_frame.pack(fill='x', pady=(0, 12), padx=pad_x)
        
        row_t = tk.Frame(title_frame, bg=Theme.BG)
        row_t.pack(fill='x', pady=2)
        
        tk.Label(row_t, text="字体:", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=6, anchor='e').pack(side='left')
        self.title_font_var = tk.StringVar()
        self._create_combobox(row_t, self.title_font_var, COMMON_FONTS_CN, width=16,
                              initial_value=self.settings.get('title', {}).get('font_cn', '方正小标宋简体')).pack(side='left', padx=3)
        
        tk.Label(row_t, text="字号:", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=5, anchor='e').pack(side='left', padx=(10, 0))
        self.title_size_var = tk.StringVar()
        self._create_combobox(row_t, self.title_size_var, [f"{name}({pt}pt)" for name, pt in FONT_SIZES], width=11,
                              initial_value=self._size_display(self.settings.get('title', {}).get('size', 22))).pack(side='left', padx=3)
        
        tk.Label(row_t, text="行距:", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=5, anchor='e').pack(side='left', padx=(10, 0))
        self.title_line_spacing_var = tk.StringVar(value=str(self.settings.get('title', {}).get('line_spacing', 29.45) or ''))
        tk.Entry(row_t, textvariable=self.title_line_spacing_var, font=get_font(11), width=5, relief='solid', bd=1).pack(side='left', padx=3)
        tk.Label(row_t, text="磅", font=get_font(10), bg=Theme.BG, fg=Theme.TEXT_MUTED).pack(side='left')
        
        # --- 一级标题 / 二级标题 ---
        self._create_section(main, "🔤 各级标题字体", pad_x)
        heading_frame = tk.Frame(main, bg=Theme.BG)
        heading_frame.pack(fill='x', pady=(0, 12), padx=pad_x)
        
        row_h1 = tk.Frame(heading_frame, bg=Theme.BG)
        row_h1.pack(fill='x', pady=2)
        tk.Label(row_h1, text="一级(一、):", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=10, anchor='e').pack(side='left')
        self.h1_font_var = tk.StringVar()
        self._create_combobox(row_h1, self.h1_font_var, COMMON_FONTS_CN, width=16,
                              initial_value=self.settings.get('heading1', {}).get('font_cn', '黑体')).pack(side='left', padx=3)
        tk.Label(row_h1, text="字号:", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=5, anchor='e').pack(side='left', padx=(10, 0))
        self.h1_size_var = tk.StringVar()
        self._create_combobox(row_h1, self.h1_size_var, [f"{name}({pt}pt)" for name, pt in FONT_SIZES], width=11,
                              initial_value=self._size_display(self.settings.get('heading1', {}).get('size', 16))).pack(side='left', padx=3)
        
        row_h2 = tk.Frame(heading_frame, bg=Theme.BG)
        row_h2.pack(fill='x', pady=2)
        tk.Label(row_h2, text="二级((一)):", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=10, anchor='e').pack(side='left')
        self.h2_font_var = tk.StringVar()
        self._create_combobox(row_h2, self.h2_font_var, COMMON_FONTS_CN, width=16,
                              initial_value=self.settings.get('heading2', {}).get('font_cn', '楷体_GB2312')).pack(side='left', padx=3)
        tk.Label(row_h2, text="字号:", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=5, anchor='e').pack(side='left', padx=(10, 0))
        self.h2_size_var = tk.StringVar()
        self._create_combobox(row_h2, self.h2_size_var, [f"{name}({pt}pt)" for name, pt in FONT_SIZES], width=11,
                              initial_value=self._size_display(self.settings.get('heading2', {}).get('size', 16))).pack(side='left', padx=3)
        
        # --- 正文格式 ---
        self._create_section(main, "📖 正文格式", pad_x)
        body_frame = tk.Frame(main, bg=Theme.BG)
        body_frame.pack(fill='x', pady=(0, 12), padx=pad_x)
        
        row_b1 = tk.Frame(body_frame, bg=Theme.BG)
        row_b1.pack(fill='x', pady=2)
        tk.Label(row_b1, text="字体:", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=6, anchor='e').pack(side='left')
        self.body_font_var = tk.StringVar()
        self._create_combobox(row_b1, self.body_font_var, COMMON_FONTS_CN, width=16,
                              initial_value=self.settings.get('body', {}).get('font_cn', '仿宋_GB2312')).pack(side='left', padx=3)
        
        tk.Label(row_b1, text="字号:", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=5, anchor='e').pack(side='left', padx=(10, 0))
        self.body_size_var = tk.StringVar()
        self._create_combobox(row_b1, self.body_size_var, [f"{name}({pt}pt)" for name, pt in FONT_SIZES], width=11,
                              initial_value=self._size_display(self.settings.get('body', {}).get('size', 16))).pack(side='left', padx=3)
        
        tk.Label(row_b1, text="行距:", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=5, anchor='e').pack(side='left', padx=(10, 0))
        self.line_spacing_var = tk.StringVar(value=str(self.settings.get('body', {}).get('line_spacing', 29.45) or ''))
        tk.Entry(row_b1, textvariable=self.line_spacing_var, font=get_font(11), width=5, relief='solid', bd=1).pack(side='left', padx=3)
        tk.Label(row_b1, text="磅", font=get_font(10), bg=Theme.BG, fg=Theme.TEXT_MUTED).pack(side='left')
        
        row_b2 = tk.Frame(body_frame, bg=Theme.BG)
        row_b2.pack(fill='x', pady=2)
        tk.Label(row_b2, text="首行缩进:", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=8, anchor='e').pack(side='left')
        self.indent_var = tk.StringVar()
        _body = self.settings.get('body', {})
        _indent = _body.get('indent', 32)
        _bsize = _body.get('size', 16) or 16
        _indent_chars = int(_indent / _bsize) if _bsize else 2
        self._create_combobox(row_b2, self.indent_var, ['0字符', '2字符', '4字符'], width=8,
                              initial_value=f'{_indent_chars}字符').pack(side='left', padx=3)
        
        tk.Label(row_b2, text="  ⓘ 正文字体/字号同时应用于: 三/四级标题、落款、附件、结束语",
                 font=get_font(9), bg=Theme.BG, fg=Theme.TEXT_MUTED).pack(side='left', padx=(10, 0))
        
        # --- 表格格式 ---
        self._create_section(main, "📊 表格格式", pad_x)
        table_frame = tk.Frame(main, bg=Theme.BG)
        table_frame.pack(fill='x', pady=(0, 12), padx=pad_x)
        
        row_tbl1 = tk.Frame(table_frame, bg=Theme.BG)
        row_tbl1.pack(fill='x', pady=2)
        
        tk.Label(row_tbl1, text="字体:", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=6, anchor='e').pack(side='left')
        self.table_font_var = tk.StringVar()
        self._create_combobox(row_tbl1, self.table_font_var, COMMON_FONTS_CN, width=16,
                              initial_value=self.settings.get('table', {}).get('font_cn', '仿宋_GB2312')).pack(side='left', padx=3)
        
        tk.Label(row_tbl1, text="字号:", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=5, anchor='e').pack(side='left', padx=(10, 0))
        self.table_size_var = tk.StringVar()
        self._create_combobox(row_tbl1, self.table_size_var, [f"{name}({pt}pt)" for name, pt in FONT_SIZES], width=11,
                              initial_value=self._size_display(self.settings.get('table', {}).get('size', 12))).pack(side='left', padx=3)
        
        tk.Label(row_tbl1, text="行距:", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=5, anchor='e').pack(side='left', padx=(10, 0))
        self.table_line_spacing_var = tk.StringVar(value=str(self.settings.get('table', {}).get('line_spacing', 22) or ''))
        tk.Entry(row_tbl1, textvariable=self.table_line_spacing_var, font=get_font(11), width=5, relief='solid', bd=1).pack(side='left', padx=3)
        tk.Label(row_tbl1, text="磅", font=get_font(10), bg=Theme.BG, fg=Theme.TEXT_MUTED).pack(side='left')
        
        row_tbl2 = tk.Frame(table_frame, bg=Theme.BG)
        row_tbl2.pack(fill='x', pady=2)
        self.table_header_bold_var = tk.BooleanVar(value=self.settings.get('table', {}).get('header_bold', True))
        tk.Checkbutton(
            row_tbl2, text="表头行加粗", variable=self.table_header_bold_var,
            font=get_font(11), bg=Theme.BG, fg=Theme.TEXT,
            activebackground=Theme.BG, selectcolor=Theme.CARD,
        ).pack(side='left', padx=(6, 0))
        
        # --- 特殊选项 ---
        self._create_section(main, "✨ 特殊选项", pad_x)
        special_frame = tk.Frame(main, bg=Theme.BG)
        special_frame.pack(fill='x', pady=(0, 12), padx=pad_x)
        
        self.first_bold_var = tk.BooleanVar(value=self.settings.get('first_line_bold', False))
        tk.Checkbutton(
            special_frame, text="正文段落首句加粗", variable=self.first_bold_var,
            font=get_font(12), bg=Theme.BG, fg=Theme.TEXT,
            activebackground=Theme.BG, selectcolor=Theme.CARD,
            padx=6, pady=3
        ).pack(anchor='w')
        
        self.page_number_var = tk.BooleanVar(value=self.settings.get('page_number', True))
        tk.Checkbutton(
            special_frame, text="添加页码", variable=self.page_number_var,
            font=get_font(12), bg=Theme.BG, fg=Theme.TEXT,
            activebackground=Theme.BG, selectcolor=Theme.CARD,
            padx=6, pady=3
        ).pack(anchor='w')
        
        # 页码字体
        pn_row = tk.Frame(special_frame, bg=Theme.BG)
        pn_row.pack(anchor='w', pady=(2, 6))
        tk.Label(pn_row, text="页码字体:", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY).pack(side='left', padx=(6, 4))
        self.page_number_font_var = tk.StringVar(value=self.settings.get('page_number_font', '宋体'))
        page_number_fonts = ['宋体', '仿宋', '仿宋_GB2312', 'Times New Roman']
        self._create_combobox(
            pn_row, self.page_number_font_var, page_number_fonts, width=16,
            initial_value=self.page_number_font_var.get()
        ).pack(side='left')
        
        # ============================================================
        #  高级设置（可折叠）
        # ============================================================
        self._create_advanced_section(main, pad_x)
        
        # ===== 底部按钮 =====
        btn_frame = tk.Frame(self, bg=Theme.BG)
        btn_frame.pack(fill='x', padx=20, pady=(10, 10))
        
        tk.Frame(btn_frame, bg=Theme.BORDER, height=1).pack(fill='x', pady=(0, 12))
        
        btn_row = tk.Frame(btn_frame, bg=Theme.BG)
        btn_row.pack(fill='x')
        
        # 恢复默认
        reset_btn = tk.Label(
            btn_row, text="恢复默认公文格式", font=get_font(11),
            bg=Theme.BG, fg=Theme.TEXT_SECONDARY, cursor='hand2'
        )
        reset_btn.pack(side='left')
        reset_btn.bind('<Button-1>', lambda e: self._reset_defaults())
        reset_btn.bind('<Enter>', lambda e: reset_btn.configure(fg=Theme.PRIMARY))
        reset_btn.bind('<Leave>', lambda e: reset_btn.configure(fg=Theme.TEXT_SECONDARY))
        
        # 保存按钮
        save_btn = tk.Frame(btn_row, bg=Theme.PRIMARY, cursor='hand2')
        save_btn.pack(side='right')
        save_label = tk.Label(
            save_btn, text="  保存设置  ", font=get_font(12, 'bold'),
            bg=Theme.PRIMARY, fg='white', pady=8, cursor='hand2'
        )
        save_label.pack()
        for w in [save_btn, save_label]:
            w.bind('<Button-1>', lambda e: self._save())
            w.bind('<Enter>', lambda e: (save_btn.configure(bg=Theme.PRIMARY_HOVER), save_label.configure(bg=Theme.PRIMARY_HOVER)))
            w.bind('<Leave>', lambda e: (save_btn.configure(bg=Theme.PRIMARY), save_label.configure(bg=Theme.PRIMARY)))
        
        cancel_btn = tk.Label(
            btn_row, text="取消", font=get_font(11),
            bg=Theme.BG, fg=Theme.TEXT_SECONDARY, cursor='hand2', padx=15
        )
        cancel_btn.pack(side='right', padx=(0, 15))
        cancel_btn.bind('<Button-1>', lambda e: self._on_close())
        
        size_grip = tk.Sizegrip(btn_frame)
        size_grip.pack(side='right', padx=(0, 2), pady=(2, 0))
    
    def _create_advanced_section(self, parent, pad_x):
        """创建可折叠的高级设置区域"""
        self._adv_expanded = False
        
        # 折叠按钮
        self._adv_toggle_frame = tk.Frame(parent, bg=Theme.BG)
        self._adv_toggle_frame.pack(fill='x', padx=pad_x, pady=(8, 0))
        
        # 分隔线
        tk.Frame(self._adv_toggle_frame, bg=Theme.BORDER, height=1).pack(fill='x', pady=(0, 8))
        
        self._adv_toggle_label = tk.Label(
            self._adv_toggle_frame,
            text="▸ 高级设置 — 按元素类型独立配置字体/行距",
            font=get_font(12, 'bold'), bg=Theme.BG, fg=Theme.TEXT_SECONDARY,
            cursor='hand2', anchor='w'
        )
        self._adv_toggle_label.pack(anchor='w')
        self._adv_toggle_label.bind('<Button-1>', lambda e: self._toggle_advanced())
        self._adv_toggle_label.bind('<Enter>', lambda e: self._adv_toggle_label.configure(fg=Theme.PRIMARY))
        self._adv_toggle_label.bind('<Leave>', lambda e: self._adv_toggle_label.configure(fg=Theme.TEXT_SECONDARY))
        
        # 高级内容区域（初始隐藏）
        self._adv_content = tk.Frame(parent, bg=Theme.BG)
        # 不 pack — 初始隐藏
        
        tk.Label(
            self._adv_content,
            text="ⓘ 此处可逐个元素类型覆盖上方快速设置的值。留空行距表示跟随正文行距。",
            font=get_font(9), bg=Theme.BG, fg=Theme.TEXT_MUTED
        ).pack(anchor='w', padx=pad_x, pady=(5, 8))
        
        # 元素类型列表
        elements = [
            ('recipient', '🏢 主送机关', '仿宋_GB2312', 16),
            ('heading1',  '1️⃣  一级标题 (一、)', '黑体', 16),
            ('heading2',  '2️⃣  二级标题 ((一))', '楷体_GB2312', 16),
            ('heading3',  '3️⃣  三级标题 (1.)', '仿宋_GB2312', 16),
            ('heading4',  '4️⃣  四级标题 ((1))', '仿宋_GB2312', 16),
            ('attachment', '📎 附件', '仿宋_GB2312', 16),
            ('closing',   '🧾 结束语', '仿宋_GB2312', 16),
            ('signature', '✒️  落款单位', '仿宋_GB2312', 16),
            ('date',      '📅 落款日期', '仿宋_GB2312', 16),
        ]
        
        for key, label, default_font, default_size in elements:
            self._create_adv_element_row(self._adv_content, pad_x, key, label, default_font, default_size)
    
    def _create_adv_element_row(self, parent, pad_x, key, label, default_font, default_size):
        """创建高级设置中的一个元素行：中文字体 + 英数字体 + 字号 + 行距"""
        row = tk.Frame(parent, bg=Theme.BG)
        row.pack(fill='x', padx=pad_x, pady=2)
        
        tk.Label(row, text=label, font=get_font(10), bg=Theme.BG, fg=Theme.TEXT, width=14, anchor='w').pack(side='left')
        
        # 中文字体
        font_var = tk.StringVar()
        self._create_combobox(row, font_var, COMMON_FONTS_CN, width=12,
                              initial_value=self.settings.get(key, {}).get('font_cn', default_font)).pack(side='left', padx=3)

        # 英数字体
        tk.Label(row, text="英数:", font=get_font(10), bg=Theme.BG, fg=Theme.TEXT_SECONDARY).pack(side='left', padx=(6, 0))
        font_en_var = tk.StringVar()
        self._create_combobox(row, font_en_var, COMMON_FONTS_EN, width=12,
                              initial_value=self.settings.get(key, {}).get('font_en', 'Times New Roman')).pack(side='left', padx=3)
        
        # 字号
        tk.Label(row, text="字号:", font=get_font(10), bg=Theme.BG, fg=Theme.TEXT_SECONDARY).pack(side='left', padx=(6, 0))
        size_var = tk.StringVar()
        self._create_combobox(row, size_var, [f"{name}({pt}pt)" for name, pt in FONT_SIZES], width=9,
                              initial_value=self._size_display(self.settings.get(key, {}).get('size', default_size))).pack(side='left', padx=3)
        
        # 行距
        tk.Label(row, text="行距:", font=get_font(10), bg=Theme.BG, fg=Theme.TEXT_SECONDARY).pack(side='left', padx=(6, 0))
        ls_val = self.settings.get(key, {}).get('line_spacing', '')
        ls_var = tk.StringVar(value=str(ls_val) if ls_val else '')
        tk.Entry(row, textvariable=ls_var, font=get_font(10), width=4, relief='solid', bd=1).pack(side='left', padx=3)
        
        # 存储变量引用
        self._adv_vars[key] = {'font': font_var, 'font_en': font_en_var, 'size': size_var, 'line_spacing': ls_var}
    
    def _toggle_advanced(self):
        """切换高级设置的折叠/展开"""
        if self._adv_expanded:
            self._adv_content.pack_forget()
            self._adv_toggle_label.configure(text="▸ 高级设置 — 按元素类型独立配置字体/行距")
            self._adv_expanded = False
        else:
            self._adv_content.pack(fill='x', after=self._adv_toggle_frame, pady=(0, 10))
            self._adv_toggle_label.configure(text="▾ 高级设置 — 按元素类型独立配置字体/行距")
            self._adv_expanded = True
        
        # 更新滚动区域
        self.content_frame.update_idletasks()
        self.canvas.configure(scrollregion=self.canvas.bbox('all'))
    
    # ==================== 滚动/辅助方法 ====================
    
    def _on_frame_configure(self, event):
        self.canvas.configure(scrollregion=self.canvas.bbox('all'))
    
    def _on_canvas_configure(self, event):
        # 保持内容宽度不小于画布宽度，允许水平滚动
        content_w = self.content_frame.winfo_reqwidth()
        self.canvas.itemconfig(self.canvas_window, width=max(event.width, content_w))
    
    def _bind_mousewheel(self):
        self.canvas.bind_all('<MouseWheel>', self._on_mousewheel)
        self.canvas.bind_all('<Button-4>', self._on_mousewheel)
        self.canvas.bind_all('<Button-5>', self._on_mousewheel)
        self.canvas.bind_all('<Shift-MouseWheel>', self._on_shift_mousewheel)
    
    def _unbind_mousewheel(self):
        self.canvas.unbind_all('<MouseWheel>')
        self.canvas.unbind_all('<Button-4>')
        self.canvas.unbind_all('<Button-5>')
        self.canvas.unbind_all('<Shift-MouseWheel>')
    
    def _on_mousewheel(self, event):
        if event.num == 4:
            self.canvas.yview_scroll(-1, 'units')
        elif event.num == 5:
            self.canvas.yview_scroll(1, 'units')
        else:
            self.canvas.yview_scroll(int(-1 * (event.delta / 120)), 'units')

    def _on_shift_mousewheel(self, event):
        if event.delta:
            self.canvas.xview_scroll(int(-1 * (event.delta / 120)), 'units')
    
    def _create_section(self, parent, title, padx=0):
        tk.Label(
            parent, text=title, font=get_font(12, 'bold'),
            bg=Theme.BG, fg=Theme.TEXT
        ).pack(anchor='w', pady=(10, 4), padx=padx)
    
    def _create_combobox(self, parent, variable, values, width=15, initial_value=None):
        """创建下拉框（OptionMenu）"""
        frame = tk.Frame(parent, bg=Theme.INPUT_BG, highlightbackground=Theme.BORDER, highlightthickness=1)
        
        if initial_value is not None:
            if initial_value in values:
                reordered = [initial_value] + [v for v in values if v != initial_value]
            else:
                reordered = [initial_value] + list(values)
        else:
            reordered = list(values)
        
        menu = tk.OptionMenu(frame, variable, *reordered)
        menu.configure(
            font=get_font(10), bg=Theme.INPUT_BG, fg=Theme.TEXT,
            activebackground=Theme.PRIMARY_LIGHT, activeforeground=Theme.TEXT,
            highlightthickness=0, relief='flat', width=width, anchor='w'
        )
        menu['menu'].configure(font=get_font(10), bg=Theme.CARD)
        menu.pack(fill='x')
        
        return frame
    
    def _size_display(self, pt_value):
        """pt值 → 显示字符串"""
        try:
            pt_value = float(pt_value)
        except (TypeError, ValueError):
            pt_value = 16.0
        for name, pt in FONT_SIZES:
            if abs(float(pt) - pt_value) < 0.01:
                return f"{name}({pt}pt)"
        return f"自定义({pt_value}pt)"
    
    def _get_size_from_var(self, var):
        """从字号下拉框获取pt值"""
        text = var.get()
        for name, pt in FONT_SIZES:
            if f"{name}({pt}pt)" == text:
                return pt
        import re as _re
        match = _re.search(r'\((\d+(?:\.\d+)?)\s*pt\)', text)
        if match:
            return float(match.group(1))
        return 16
    
    def _get_line_spacing(self, var, fallback=29.45):
        """从行距输入框获取值，空值返回 fallback"""
        val = var.get().strip()
        if not val:
            return fallback
        try:
            return int(float(val))
        except ValueError:
            return fallback
    
    # ==================== 加载/保存 ====================
    
    def _load_values(self):
        """加载设置到 UI"""
        s = self.settings
        try:
            # 页边距
            for key in ['top', 'bottom', 'left', 'right']:
                self.margin_vars[key].set(str(s.get('page', {}).get(key, 2.5)))
            
            # 标题
            self.title_font_var.set(s.get('title', {}).get('font_cn', '方正小标宋简体'))
            self._set_size_var(self.title_size_var, s.get('title', {}).get('size', 22))
            self.title_line_spacing_var.set(str(s.get('title', {}).get('line_spacing', 29.45) or ''))
            
            # 一/二级标题
            self.h1_font_var.set(s.get('heading1', {}).get('font_cn', '黑体'))
            self._set_size_var(self.h1_size_var, s.get('heading1', {}).get('size', 16))
            self.h2_font_var.set(s.get('heading2', {}).get('font_cn', '楷体_GB2312'))
            self._set_size_var(self.h2_size_var, s.get('heading2', {}).get('size', 16))
            
            # 正文
            self.body_font_var.set(s.get('body', {}).get('font_cn', '仿宋_GB2312'))
            self._set_size_var(self.body_size_var, s.get('body', {}).get('size', 16))
            self.line_spacing_var.set(str(s.get('body', {}).get('line_spacing', 29.45) or ''))
            
            body_size = s.get('body', {}).get('size', 16) or 16
            indent = s.get('body', {}).get('indent', 32)
            indent_chars = int(indent / body_size) if body_size else 2
            self.indent_var.set(f'{indent_chars}字符')
            
            # 表格
            tbl = s.get('table', {})
            self.table_font_var.set(tbl.get('font_cn', '仿宋_GB2312'))
            self._set_size_var(self.table_size_var, tbl.get('size', 12))
            self.table_line_spacing_var.set(str(tbl.get('line_spacing', 22) or ''))
            self.table_header_bold_var.set(tbl.get('header_bold', True))
            
            # 特殊选项
            self.first_bold_var.set(s.get('first_line_bold', False))
            self.page_number_var.set(s.get('page_number', True))
            self.page_number_font_var.set(s.get('page_number_font', '宋体'))
            
            # 高级设置
            for key, vars_dict in self._adv_vars.items():
                elem = s.get(key, {})
                vars_dict['font'].set(elem.get('font_cn', '仿宋_GB2312'))
                vars_dict['font_en'].set(elem.get('font_en', 'Times New Roman'))
                self._set_size_var(vars_dict['size'], elem.get('size', 16))
                ls = elem.get('line_spacing', '')
                vars_dict['line_spacing'].set(str(ls) if ls else '')
        except Exception as e:
            print(f"[警告] 加载设置到界面失败: {e}")
    
    def _set_size_var(self, var, pt_value):
        try:
            pt_value = float(pt_value)
        except (TypeError, ValueError):
            pt_value = 16.0
        for name, pt in FONT_SIZES:
            if abs(float(pt) - pt_value) < 0.01:
                var.set(f"{name}({pt}pt)")
                return
        var.set(f"自定义({pt_value}pt)")
    
    def _reset_defaults(self):
        import copy
        self.settings = copy.deepcopy(DEFAULT_CUSTOM_SETTINGS)
        self._load_values()
    
    def _save(self):
        """保存设置 - 快速设置为主，高级设置覆盖"""
        try:
            # 收集快速设置值
            page = {key: float(self.margin_vars[key].get()) for key in ['top', 'bottom', 'left', 'right']}
            
            title_size = self._get_size_from_var(self.title_size_var)
            h1_size = self._get_size_from_var(self.h1_size_var)
            h2_size = self._get_size_from_var(self.h2_size_var)
            body_size = self._get_size_from_var(self.body_size_var)
            body_ls = self._get_line_spacing(self.line_spacing_var, 29.45)
            title_ls = self._get_line_spacing(self.title_line_spacing_var, 29.45)
            
            # 首行缩进
            indent_text = self.indent_var.get()
            indent_chars = int(indent_text.replace('字符', ''))
            indent_pt = indent_chars * body_size
            
            body_font = self.body_font_var.get()
            
            # 构建基础设置 — 正文字体联动到多个元素
            self.settings = {
                'name': '自定义格式',
                'page': page,
                'title': {
                    'font_cn': self.title_font_var.get(), 'font_en': 'Times New Roman',
                    'size': title_size, 'bold': False, 'align': 'center', 'indent': 0,
                    'line_spacing': title_ls, 'space_before': 0, 'space_after': 0
                },
                'recipient': {
                    'font_cn': body_font, 'font_en': 'Times New Roman',
                    'size': body_size, 'bold': False, 'align': 'left', 'indent': 0,
                    'line_spacing': body_ls, 'space_before': 0, 'space_after': 0
                },
                'heading1': {
                    'font_cn': self.h1_font_var.get(), 'font_en': 'Times New Roman',
                    'size': h1_size, 'bold': False, 'align': 'left', 'indent': indent_pt,
                    'line_spacing': body_ls, 'space_before': 0, 'space_after': 0
                },
                'heading2': {
                    'font_cn': self.h2_font_var.get(), 'font_en': 'Times New Roman',
                    'size': h2_size, 'bold': False, 'align': 'left', 'indent': indent_pt,
                    'line_spacing': body_ls, 'space_before': 0, 'space_after': 0
                },
                'heading3': {
                    'font_cn': body_font, 'font_en': 'Times New Roman',
                    'size': body_size, 'bold': False, 'align': 'left', 'indent': indent_pt,
                    'line_spacing': body_ls, 'space_before': 0, 'space_after': 0
                },
                'heading4': {
                    'font_cn': body_font, 'font_en': 'Times New Roman',
                    'size': body_size, 'bold': False, 'align': 'left', 'indent': indent_pt,
                    'line_spacing': body_ls, 'space_before': 0, 'space_after': 0
                },
                'body': {
                    'font_cn': body_font, 'font_en': 'Times New Roman',
                    'size': body_size, 'bold': False, 'align': 'justify', 'indent': indent_pt,
                    'line_spacing': body_ls, 'space_before': 0, 'space_after': 0
                },
                'signature': {
                    'font_cn': body_font, 'font_en': 'Times New Roman',
                    'size': body_size, 'bold': False, 'align': 'right', 'indent': 0,
                    'line_spacing': body_ls, 'space_before': 0, 'space_after': 0
                },
                'date': {
                    'font_cn': body_font, 'font_en': 'Times New Roman',
                    'size': body_size, 'bold': False, 'align': 'right', 'indent': 0,
                    'line_spacing': body_ls, 'space_before': 0, 'space_after': 0
                },
                'attachment': {
                    'font_cn': body_font, 'font_en': 'Times New Roman',
                    'size': body_size, 'bold': False, 'align': 'justify', 'indent': indent_pt,
                    'line_spacing': body_ls, 'space_before': 0, 'space_after': 0
                },
                'closing': {
                    'font_cn': body_font, 'font_en': 'Times New Roman',
                    'size': body_size, 'bold': False, 'align': 'left', 'indent': indent_pt,
                    'line_spacing': body_ls, 'space_before': 0, 'space_after': 0
                },
                'table': {
                    'font_cn': self.table_font_var.get(), 'font_en': 'Times New Roman',
                    'size': self._get_size_from_var(self.table_size_var), 'bold': False,
                    'line_spacing': self._get_line_spacing(self.table_line_spacing_var, 22),
                    'first_line_indent': 0,
                    'header_bold': self.table_header_bold_var.get()
                },
                'first_line_bold': self.first_bold_var.get(),
                'page_number': self.page_number_var.get(),
                'page_number_font': self.page_number_font_var.get()
            }
            
            # 应用高级设置覆盖（如果用户有修改）
            for key, vars_dict in self._adv_vars.items():
                if key in self.settings and isinstance(self.settings[key], dict):
                    adv_font = vars_dict['font'].get()
                    adv_font_en = vars_dict['font_en'].get()
                    adv_size = self._get_size_from_var(vars_dict['size'])
                    adv_ls_str = vars_dict['line_spacing'].get().strip()
                    
                    # 只在高级值与快速设置不同的时候覆盖
                    if adv_font:
                        self.settings[key]['font_cn'] = adv_font
                    if adv_font_en:
                        self.settings[key]['font_en'] = adv_font_en
                    if adv_size:
                        self.settings[key]['size'] = adv_size
                    if adv_ls_str:
                        try:
                            self.settings[key]['line_spacing'] = int(float(adv_ls_str))
                        except ValueError:
                            pass
            
            save_custom_settings(self.settings)
            
            if self.on_save:
                self.on_save(self.settings)
            
            messagebox.showinfo("保存成功", "自定义格式设置已保存", parent=self)
            self.destroy()
            
        except ValueError as e:
            messagebox.showerror("输入错误", f"请检查输入的数值是否正确：\n{e}", parent=self)
    
    def _on_close(self):
        result = messagebox.askyesnocancel("保存设置", "是否保存当前设置？", parent=self)
        if result is None:
            return
        if result:
            self._save()
        else:
            self.destroy()


# ===== 大尺寸线条图标 =====
class Icons:
    """用 Canvas 绘制的线条图标 - 48px 大尺寸"""
    
    @staticmethod
    def draw_magic(canvas, x, y, size=48, color='#2E2E2E'):
        """智能处理 - 魔法棒"""
        s = size
        lw = 2.5  # 线宽
        # 魔法棒主体
        canvas.create_line(x+s*0.15, y+s*0.85, x+s*0.65, y+s*0.35, fill=color, width=lw, capstyle='round')
        # 星星点缀
        stars = [(0.7, 0.2), (0.85, 0.35), (0.75, 0.5), (0.55, 0.15)]
        for px, py in stars:
            r = 3
            canvas.create_oval(x+s*px-r, y+s*py-r, x+s*px+r, y+s*py+r, fill=color, outline='')
        # 光芒线
        canvas.create_line(x+s*0.7, y+s*0.08, x+s*0.7, y+s*0.22, fill=color, width=1.5)
        canvas.create_line(x+s*0.9, y+s*0.28, x+s*0.78, y+s*0.35, fill=color, width=1.5)
    
    @staticmethod
    def draw_search(canvas, x, y, size=48, color='#2E2E2E'):
        """诊断 - 放大镜"""
        s = size
        lw = 2.5
        # 镜框
        canvas.create_oval(x+s*0.12, y+s*0.12, x+s*0.58, y+s*0.58, outline=color, width=lw)
        # 镜柄
        canvas.create_line(x+s*0.52, y+s*0.52, x+s*0.85, y+s*0.85, fill=color, width=lw, capstyle='round')
        # 高光
        canvas.create_arc(x+s*0.18, y+s*0.18, x+s*0.4, y+s*0.4, start=120, extent=60, style='arc', outline=color, width=1.5)
    
    @staticmethod
    def draw_edit(canvas, x, y, size=48, color='#2E2E2E'):
        """标点修复 - 铅笔"""
        s = size
        lw = 2.5
        # 笔身
        canvas.create_line(x+s*0.2, y+s*0.8, x+s*0.7, y+s*0.3, fill=color, width=lw, capstyle='round')
        # 笔尖
        canvas.create_polygon(
            x+s*0.15, y+s*0.85,
            x+s*0.2, y+s*0.8,
            x+s*0.25, y+s*0.85,
            fill=color, outline=''
        )
        # 笔头
        canvas.create_line(x+s*0.7, y+s*0.3, x+s*0.8, y+s*0.2, fill=color, width=lw, capstyle='round')
        canvas.create_line(x+s*0.75, y+s*0.35, x+s*0.85, y+s*0.25, fill=color, width=lw, capstyle='round')
    
    @staticmethod
    def draw_file(canvas, x, y, size=48, color='#2E2E2E'):
        """文件图标"""
        s = size
        lw = 2
        # 文件主体
        points = [
            x+s*0.2, y+s*0.1,   # 左上
            x+s*0.2, y+s*0.9,   # 左下
            x+s*0.8, y+s*0.9,   # 右下
            x+s*0.8, y+s*0.3,   # 右上（折角下）
            x+s*0.6, y+s*0.1,   # 折角
        ]
        canvas.create_polygon(points, fill='', outline=color, width=lw)
        # 折角线
        canvas.create_line(x+s*0.6, y+s*0.1, x+s*0.6, y+s*0.3, fill=color, width=lw)
        canvas.create_line(x+s*0.6, y+s*0.3, x+s*0.8, y+s*0.3, fill=color, width=lw)
    
    @staticmethod
    def draw_check(canvas, x, y, size=32, color='#7CB87C'):
        """勾选"""
        s = size
        canvas.create_line(x+s*0.15, y+s*0.5, x+s*0.4, y+s*0.75, fill=color, width=3, capstyle='round')
        canvas.create_line(x+s*0.4, y+s*0.75, x+s*0.85, y+s*0.25, fill=color, width=3, capstyle='round')


class FileInputField(tk.Frame):
    """文件输入框 - 带明显容器"""
    
    def __init__(self, parent, label_text, placeholder, variable, command, **kwargs):
        super().__init__(parent, bg=Theme.BG, **kwargs)
        
        self.variable = variable
        self.command = command
        self.placeholder = placeholder
        
        # 标签
        tk.Label(
            self,
            text=label_text,
            font=get_font(11),
            bg=Theme.BG,
            fg=Theme.TEXT_SECONDARY,
            width=4,
            anchor='w'
        ).pack(side='left')
        
        # 输入框容器
        self.container = tk.Frame(
            self,
            bg=Theme.INPUT_BG,
            highlightbackground=Theme.BORDER,
            highlightcolor=Theme.PRIMARY,
            highlightthickness=1
        )
        self.container.pack(side='left', fill='x', expand=True, padx=(Theme.SPACE_SM, 0))
        
        inner = tk.Frame(self.container, bg=Theme.INPUT_BG)
        inner.pack(fill='both', expand=True, padx=Theme.SPACE_MD, pady=Theme.SPACE_SM + 2)
        
        # 文件名显示
        self.filename_label = tk.Label(
            inner,
            text="未选择",
            font=get_font(11),
            bg=Theme.INPUT_BG,
            fg=Theme.TEXT_MUTED,
            anchor='w'
        )
        self.filename_label.pack(side='left', fill='x', expand=True)
        
        # 分隔线
        tk.Frame(inner, bg=Theme.BORDER, width=1).pack(side='left', fill='y', padx=Theme.SPACE_MD)
        
        # 操作按钮
        self.action_btn = tk.Label(
            inner,
            text=placeholder,
            font=get_font(10),
            bg=Theme.INPUT_BG,
            fg=Theme.PRIMARY,
            cursor='hand2'
        )
        self.action_btn.pack(side='right')
        
        # 绑定点击
        for widget in [self.container, inner, self.filename_label, self.action_btn]:
            widget.bind('<Button-1>', self._on_click)
            widget.configure(cursor='hand2')
        
        # 悬停效果
        self.container.bind('<Enter>', lambda e: self.container.configure(highlightbackground='#D0CCC6'))
        self.container.bind('<Leave>', lambda e: self.container.configure(highlightbackground=Theme.BORDER))
        
        # 监听变量
        self.variable.trace_add('write', self._update_display)
    
    def _on_click(self, event=None):
        if self.command:
            self.command()
    
    def _update_display(self, *args):
        path = self.variable.get()
        if path:
            # 显示文件名，路径过长则截断
            filename = Path(path).name
            if len(filename) > 40:
                filename = filename[:37] + "..."
            self.filename_label.configure(text=filename, fg=Theme.TEXT)
        else:
            self.filename_label.configure(text="未选择", fg=Theme.TEXT_MUTED)


class SelectableCard(tk.Frame):
    """可选择的卡片 - 大图标版"""
    
    def __init__(self, parent, title, description, value, variable,
                 icon_draw_func=None, is_featured=False, command=None, **kwargs):
        
        bg_color = Theme.CARD_ALT if is_featured else Theme.CARD
        super().__init__(parent, bg=bg_color, **kwargs)
        
        self.value = value
        self.variable = variable
        self.command = command
        self.is_featured = is_featured
        self.bg_color = bg_color
        self.selected = False
        
        # 边框
        self.configure(
            highlightbackground=Theme.BORDER,
            highlightcolor=Theme.BORDER_SELECTED,
            highlightthickness=1
        )
        
        # 内容 - 水平布局：左图标 + 右文字
        content = tk.Frame(self, bg=bg_color)
        content.pack(fill='both', expand=True, padx=Theme.SPACE_LG, pady=Theme.SPACE_LG)
        
        # 左侧：图标
        if icon_draw_func:
            icon_size = 56 if is_featured else 48
            self.icon_canvas = tk.Canvas(
                content,
                width=icon_size + 8,
                height=icon_size + 8,
                bg=bg_color,
                highlightthickness=0
            )
            self.icon_canvas.pack(side='left', padx=(0, Theme.SPACE_MD))
            icon_draw_func(self.icon_canvas, 4, 4, icon_size, Theme.TEXT)
            self._bind_click(self.icon_canvas)
        
        # 右侧：文字区域
        text_frame = tk.Frame(content, bg=bg_color)
        text_frame.pack(side='left', fill='both', expand=True)
        
        # 标题行（标题 + 推荐标签）
        title_row = tk.Frame(text_frame, bg=bg_color)
        title_row.pack(fill='x', anchor='w')
        
        title_size = 16 if is_featured else 14
        self.title_label = tk.Label(
            title_row,
            text=title,
            font=get_font(title_size, 'bold'),
            bg=bg_color,
            fg=Theme.TEXT,
            anchor='w'
        )
        self.title_label.pack(side='left')
        
        # 推荐标签
        if is_featured:
            tag = tk.Label(
                title_row,
                text=" 推荐 ",
                font=get_font(10, 'bold'),
                bg=Theme.PRIMARY,
                fg='white',
                padx=10,
                pady=3
            )
            tag.pack(side='left', padx=(Theme.SPACE_SM, 0))
            self._bind_click(tag)
        
        self._bind_click(title_row)
        
        # 描述
        desc_size = 12 if is_featured else 11
        self.desc_label = tk.Label(
            text_frame,
            text=description,
            font=get_font(desc_size),
            bg=bg_color,
            fg=Theme.TEXT_SECONDARY,
            anchor='w',
            justify='left'
        )
        self.desc_label.pack(fill='x', anchor='w', pady=(Theme.SPACE_SM, 0))
        
        # 绑定事件
        self._bind_click(self)
        self._bind_click(content)
        self._bind_click(text_frame)
        self._bind_click(self.title_label)
        self._bind_click(self.desc_label)
        
        # 监听变量
        self.variable.trace_add('write', self._on_variable_change)
        self._update_style()
    
    def _bind_click(self, widget):
        widget.bind('<Button-1>', self._on_click)
        widget.bind('<Enter>', self._on_enter)
        widget.bind('<Leave>', self._on_leave)
        widget.configure(cursor='hand2')
    
    def _on_click(self, event=None):
        self.variable.set(self.value)
        if self.command:
            self.command()
    
    def _on_enter(self, event=None):
        if not self.selected:
            self.configure(highlightbackground='#D0CCC6')
    
    def _on_leave(self, event=None):
        self._update_style()
    
    def _on_variable_change(self, *args):
        self._update_style()
    
    def _update_style(self):
        self.selected = (self.variable.get() == self.value)
        if self.selected:
            self.configure(highlightbackground=Theme.BORDER_SELECTED, highlightthickness=2)
        else:
            self.configure(highlightbackground=Theme.BORDER, highlightthickness=1)


class PresetCard(tk.Frame):
    """格式预设卡片"""
    
    def __init__(self, parent, text, value, variable, command=None, **kwargs):
        super().__init__(parent, bg=Theme.CARD, **kwargs)
        
        self.value = value
        self.variable = variable
        self.selected = False
        self.command = command  # 自定义点击命令
        
        self.configure(
            highlightbackground=Theme.BORDER,
            highlightcolor=Theme.BORDER_SELECTED,
            highlightthickness=1
        )
        
        self.label = tk.Label(
            self,
            text=text,
            font=get_font(12),
            bg=Theme.CARD,
            fg=Theme.TEXT,
            padx=Theme.SPACE_LG,
            pady=Theme.SPACE_MD
        )
        self.label.pack()
        
        # 绑定
        for widget in [self, self.label]:
            widget.bind('<Button-1>', self._on_click)
            widget.bind('<Enter>', self._on_enter)
            widget.bind('<Leave>', self._on_leave)
            widget.configure(cursor='hand2')
        
        self.variable.trace_add('write', self._update_style)
        self._update_style()
    
    def _on_click(self, event=None):
        self.variable.set(self.value)
        # 如果有自定义命令，执行它
        if self.command:
            self.command()
    
    def _on_enter(self, event=None):
        if not self.selected:
            self.configure(highlightbackground='#D0CCC6')
    
    def _on_leave(self, event=None):
        self._update_style()
    
    def _update_style(self, *args):
        self.selected = (self.variable.get() == self.value)
        if self.selected:
            self.configure(bg=Theme.PRIMARY_LIGHT, highlightbackground=Theme.PRIMARY, highlightthickness=2)
            self.label.configure(bg=Theme.PRIMARY_LIGHT, fg=Theme.TEXT, font=get_font(12, 'bold'))
        else:
            self.configure(bg=Theme.CARD, highlightbackground=Theme.BORDER, highlightthickness=1)
            self.label.configure(bg=Theme.CARD, fg=Theme.TEXT, font=get_font(12))
    
    def set_enabled(self, enabled):
        if enabled:
            self.label.configure(fg=Theme.TEXT, cursor='hand2')
            self.configure(cursor='hand2')
        else:
            self.label.configure(fg=Theme.TEXT_MUTED, cursor='arrow')
            self.configure(cursor='arrow', highlightbackground=Theme.BORDER_LIGHT)


class CollapsibleLog(tk.Frame):
    """可折叠的日志区域"""
    
    def __init__(self, parent, **kwargs):
        super().__init__(parent, bg=Theme.BG, **kwargs)
        
        self.expanded = False
        
        # 折叠条
        self.toggle_bar = tk.Frame(self, bg='#E8E4DE', height=36)
        self.toggle_bar.pack(fill='x')
        self.toggle_bar.pack_propagate(False)
        
        self.toggle_btn = tk.Label(
            self.toggle_bar,
            text="＋  展开运行日志",
            font=get_font(11),
            bg='#E8E4DE',
            fg=Theme.TEXT_SECONDARY,
            cursor='hand2'
        )
        self.toggle_btn.pack(side='left', padx=Theme.SPACE_MD, pady=Theme.SPACE_SM)
        self.toggle_btn.bind('<Button-1>', self._toggle)
        self.toggle_bar.bind('<Button-1>', self._toggle)
        self.toggle_bar.configure(cursor='hand2')
        
        # 日志面板
        self.log_panel = tk.Frame(self, bg=Theme.LOG_BG)
        
        # 日志文本
        self.log_text = tk.Text(
            self.log_panel,
            font=('Consolas', 11),
            bg=Theme.LOG_BG,
            fg=Theme.LOG_TEXT,
            relief='flat',
            padx=Theme.SPACE_LG,
            pady=Theme.SPACE_MD,
            wrap='word',
            height=10,
            highlightthickness=0,
            insertbackground=Theme.LOG_TEXT
        )
        self.log_text.pack(side='left', fill='both', expand=True)
        
        # 配置颜色标签
        self.log_text.tag_configure('info', foreground=Theme.LOG_TEXT)
        self.log_text.tag_configure('success', foreground=Theme.LOG_SUCCESS)
        self.log_text.tag_configure('warning', foreground=Theme.LOG_WARNING)
        self.log_text.tag_configure('error', foreground=Theme.LOG_ERROR)
    
    def _toggle(self, event=None):
        self.expanded = not self.expanded
        if self.expanded:
            self.log_panel.pack(fill='both', expand=True)
            self.toggle_btn.configure(text="－  收起运行日志")
        else:
            self.log_panel.pack_forget()
            self.toggle_btn.configure(text="＋  展开运行日志")
    
    def log(self, message, tag='info'):
        self.log_text.insert(tk.END, message + "\n", tag)
        self.log_text.see(tk.END)
    
    def clear(self):
        self.log_text.delete(1.0, tk.END)


class ResultPanel(tk.Frame):
    """结果反馈面板"""
    
    def __init__(self, parent, **kwargs):
        super().__init__(parent, bg=Theme.BG, **kwargs)
        
        # 占位状态
        self.placeholder = tk.Label(
            self,
            text="处理结果将在此处显示",
            font=get_font(12),
            bg=Theme.BG,
            fg=Theme.TEXT_MUTED,
            pady=Theme.SPACE_XL
        )
        self.placeholder.pack()
        
        # 结果卡片
        self.result_card = tk.Frame(self, bg=Theme.CARD, highlightbackground=Theme.BORDER, highlightthickness=1)
        self.result_content = tk.Frame(self.result_card, bg=Theme.CARD)
        self.result_content.pack(fill='both', expand=True, padx=Theme.SPACE_LG, pady=Theme.SPACE_LG)
    
    def show_success(self, message, filepath=None):
        self.placeholder.pack_forget()
        
        for widget in self.result_content.winfo_children():
            widget.destroy()
        
        # 成功图标 + 消息
        header = tk.Frame(self.result_content, bg=Theme.CARD)
        header.pack(fill='x', anchor='w')
        
        icon_canvas = tk.Canvas(header, width=36, height=36, bg=Theme.CARD, highlightthickness=0)
        icon_canvas.pack(side='left')
        Icons.draw_check(icon_canvas, 2, 2, 32, Theme.LOG_SUCCESS)
        
        tk.Label(
            header,
            text=message,
            font=get_font(15, 'bold'),
            bg=Theme.CARD,
            fg=Theme.TEXT,
            anchor='w'
        ).pack(side='left', padx=(Theme.SPACE_SM, 0))
        
        if filepath:
            tk.Label(
                self.result_content,
                text=f"输出文件：{filepath}",
                font=get_font(11),
                bg=Theme.CARD,
                fg=Theme.TEXT_SECONDARY,
                anchor='w'
            ).pack(fill='x', anchor='w', pady=(Theme.SPACE_SM, 0))
        
        self.result_card.pack(fill='x', pady=(Theme.SPACE_MD, 0))
    
    def show_diagnosis(self, results):
        self.placeholder.pack_forget()
        
        for widget in self.result_content.winfo_children():
            widget.destroy()
        
        tk.Label(
            self.result_content,
            text="诊断报告",
            font=get_font(15, 'bold'),
            bg=Theme.CARD,
            fg=Theme.TEXT,
            anchor='w'
        ).pack(fill='x', anchor='w', pady=(0, Theme.SPACE_MD))
        
        total = 0
        categories = [
            ('标点问题', results.get('punctuation', [])),
            ('序号问题', results.get('numbering', [])),
            ('段落问题', results.get('paragraph', [])),
            ('字体问题', results.get('font', [])),
        ]
        
        for name, issues in categories:
            count = len(issues)
            total += count
            
            row = tk.Frame(self.result_content, bg=Theme.CARD)
            row.pack(fill='x', pady=3)
            
            tk.Label(
                row,
                text=name,
                font=get_font(12),
                bg=Theme.CARD,
                fg=Theme.TEXT,
                width=10,
                anchor='w'
            ).pack(side='left')
            
            count_color = Theme.LOG_WARNING if count > 0 else Theme.LOG_SUCCESS
            tk.Label(
                row,
                text=f"{count} 处" if count > 0 else "无问题",
                font=get_font(12),
                bg=Theme.CARD,
                fg=count_color,
                anchor='w'
            ).pack(side='left')
        
        tk.Frame(self.result_content, bg=Theme.BORDER, height=1).pack(fill='x', pady=Theme.SPACE_MD)
        
        summary_color = Theme.LOG_SUCCESS if total == 0 else Theme.LOG_WARNING
        summary_text = "文档格式规范，未发现问题" if total == 0 else f"共发现 {total} 处格式问题"
        
        tk.Label(
            self.result_content,
            text=summary_text,
            font=get_font(13, 'bold'),
            bg=Theme.CARD,
            fg=summary_color,
            anchor='w'
        ).pack(fill='x', anchor='w')
        
        self.result_card.pack(fill='x', pady=(Theme.SPACE_MD, 0))
    
    def reset(self):
        self.result_card.pack_forget()
        for widget in self.result_content.winfo_children():
            widget.destroy()
        self.placeholder.pack()


class DocFormatApp:
    def __init__(self, root):
        self.root = root
        self.root.title("公文格式处理工具")
        self.root.geometry("750x900")
        self.root.minsize(680, 750)
        self.root.configure(bg=Theme.BG)
        
        # 变量
        self.input_file = tk.StringVar()
        self.output_file = tk.StringVar()
        self.operation = tk.StringVar(value="smart")
        self.preset = tk.StringVar(value="official")
        
        self.preset_cards = []
        
        self.create_widgets()
    
    def create_widgets(self):
        """构建界面"""
        # 主容器 - 带滚动
        container = tk.Frame(self.root, bg=Theme.BG)
        container.pack(fill='both', expand=True)
        
        # Canvas + 自定义滚动条
        self.canvas = tk.Canvas(container, bg=Theme.BG, highlightthickness=0)
        self.scrollbar_canvas = tk.Canvas(container, width=14, bg=Theme.BG, highlightthickness=0)
        
        self.canvas.pack(side='left', fill='both', expand=True)
        self.scrollbar_canvas.pack(side='right', fill='y')
        
        # 内容Frame
        self.main_frame = tk.Frame(self.canvas, bg=Theme.BG)
        self.canvas_window = self.canvas.create_window((0, 0), window=self.main_frame, anchor='nw')
        
        # 绑定滚动
        self.main_frame.bind('<Configure>', self._on_frame_configure)
        self.canvas.bind('<Configure>', self._on_canvas_configure)
        self.root.bind_all('<MouseWheel>', self._on_mousewheel)
        self.scrollbar_canvas.bind('<Button-1>', self._on_scrollbar_click)
        self.scrollbar_canvas.bind('<B1-Motion>', self._on_scrollbar_drag)
        
        # 内容区域
        content = tk.Frame(self.main_frame, bg=Theme.BG)
        content.pack(fill='both', expand=True, padx=Theme.SPACE_XL, pady=Theme.SPACE_LG)
        
        # ===== 1. 头部 =====
        tk.Label(
            content,
            text="公文格式处理工具",
            font=get_font(24, 'bold'),
            bg=Theme.BG,
            fg=Theme.TEXT
        ).pack(anchor='w', pady=(0, Theme.SPACE_XL))
        
        # ===== 2. 文件选择区 =====
        file_section = tk.Frame(content, bg=Theme.BG)
        file_section.pack(fill='x', pady=(0, Theme.SPACE_LG))
        
        self.input_field = FileInputField(
            file_section,
            label_text="输入",
            placeholder="点击选择需要修改的文档",
            variable=self.input_file,
            command=self.browse_input
        )
        self.input_field.pack(fill='x', pady=(0, Theme.SPACE_SM))
        
        self.output_field = FileInputField(
            file_section,
            label_text="输出",
            placeholder="文档修改后的储存位置",
            variable=self.output_file,
            command=self.browse_output
        )
        self.output_field.pack(fill='x')
        
        # 分隔
        tk.Frame(content, bg=Theme.BORDER, height=1).pack(fill='x', pady=Theme.SPACE_LG)
        
        # ===== 3. 功能选择区 =====
        mode_section = tk.Frame(content, bg=Theme.BG)
        mode_section.pack(fill='x', pady=(0, Theme.SPACE_LG))
        
        # 大卡片 - 智能一键处理
        smart_card = SelectableCard(
            mode_section,
            title="智能一键处理",
            description="自动修复标点符号，并应用标准格式规范，一步到位完成文档处理",
            value="smart",
            variable=self.operation,
            icon_draw_func=Icons.draw_magic,
            is_featured=True,
            command=self._on_mode_change
        )
        smart_card.pack(fill='x', pady=(0, Theme.SPACE_MD))
        
        # 两个小卡片
        small_cards = tk.Frame(mode_section, bg=Theme.BG)
        small_cards.pack(fill='x')
        small_cards.columnconfigure(0, weight=1)
        small_cards.columnconfigure(1, weight=1)
        
        diag_card = SelectableCard(
            small_cards,
            title="格式诊断",
            description="仅分析文档问题，不修改文件",
            value="analyze",
            variable=self.operation,
            icon_draw_func=Icons.draw_search,
            command=self._on_mode_change
        )
        diag_card.grid(row=0, column=0, sticky='nsew', padx=(0, Theme.SPACE_SM))
        
        punct_card = SelectableCard(
            small_cards,
            title="标点修复",
            description="仅修复中英文标点混用",
            value="punctuation",
            variable=self.operation,
            icon_draw_func=Icons.draw_edit,
            command=self._on_mode_change
        )
        punct_card.grid(row=0, column=1, sticky='nsew')
        
        # ===== 4. 格式预设 =====
        preset_section = tk.Frame(content, bg=Theme.BG)
        preset_section.pack(fill='x', pady=(0, Theme.SPACE_LG))
        
        # 标题行
        tk.Label(
            preset_section,
            text="格式预设",
            font=get_font(12),
            bg=Theme.BG,
            fg=Theme.TEXT_SECONDARY
        ).pack(anchor='w', pady=(0, Theme.SPACE_SM))
        
        preset_row = tk.Frame(preset_section, bg=Theme.BG)
        preset_row.pack(fill='x')
        
        presets = [
            ('official', 'GB/T 公文标准'),
            ('academic', '学术论文'),
            ('legal', '法律文书'),
        ]
        
        for i, (value, text) in enumerate(presets):
            card = PresetCard(preset_row, text, value, self.preset)
            card.pack(side='left', padx=(0 if i == 0 else Theme.SPACE_SM, 0))
            self.preset_cards.append(card)
        
        # 自定义卡片 - 点击直接打开设置窗口
        self.custom_card = PresetCard(
            preset_row, '⚙️ 自定义', 'custom', self.preset,
            command=self._open_custom_settings  # 点击时打开设置窗口
        )
        self.custom_card.pack(side='left', padx=(Theme.SPACE_SM, 0))
        self.preset_cards.append(self.custom_card)
        
        
        # ===== 5. 执行按钮 =====
        self.run_btn = tk.Frame(content, bg=Theme.PRIMARY, cursor='hand2')
        self.run_btn.pack(fill='x', pady=Theme.SPACE_LG)
        
        self.run_label = tk.Label(
            self.run_btn,
            text="开始处理",
            font=get_font(15, 'bold'),
            bg=Theme.PRIMARY,
            fg='white',
            pady=Theme.SPACE_MD + 2
        )
        self.run_label.pack()
        
        for widget in [self.run_btn, self.run_label]:
            widget.bind('<Button-1>', lambda e: self.run_operation())
            widget.bind('<Enter>', lambda e: self._btn_hover(True))
            widget.bind('<Leave>', lambda e: self._btn_hover(False))
        self.run_label.configure(cursor='hand2')
        
        # ===== 6. 结果反馈区 =====
        self.result_panel = ResultPanel(content)
        self.result_panel.pack(fill='x', pady=(0, Theme.SPACE_LG))
        
        # ===== 7. 日志区 =====
        self.log_panel = CollapsibleLog(content)
        self.log_panel.pack(fill='x', pady=(Theme.SPACE_MD, 0))
        
        # 初始化
        self._on_mode_change()
        self.log_panel.log("工具已就绪，请选择文件", 'info')
    
    def _on_frame_configure(self, event):
        self.canvas.configure(scrollregion=self.canvas.bbox('all'))
        self._draw_scrollbar()
    
    def _on_canvas_configure(self, event):
        self.canvas.itemconfig(self.canvas_window, width=event.width)
        self._draw_scrollbar()
    
    def _on_mousewheel(self, event):
        self.canvas.yview_scroll(int(-1 * (event.delta / 120)), 'units')
        self._draw_scrollbar()
    
    def _draw_scrollbar(self):
        """绘制自定义滚动条"""
        self.scrollbar_canvas.delete('all')
        
        try:
            top, bottom = self.canvas.yview()
        except:
            return
        
        if bottom - top >= 0.99:
            return
        
        w = 14
        h = self.scrollbar_canvas.winfo_height()
        
        if h < 10:
            return
        
        bar_h = max(40, (bottom - top) * h)
        bar_y = top * (h - bar_h)
        
        # 轨道
        self.scrollbar_canvas.create_rectangle(
            4, 8, w - 4, h - 8,
            fill='#E8E4DE', outline=''
        )
        
        # 滑块（更深的颜色）
        self.scrollbar_canvas.create_rectangle(
            4, bar_y + 8, w - 4, bar_y + bar_h - 8,
            fill='#A09890', outline=''
        )
    
    def _on_scrollbar_click(self, event):
        """滚动条点击"""
        try:
            h = self.scrollbar_canvas.winfo_height()
            fraction = event.y / h
            self.canvas.yview_moveto(fraction)
            self._draw_scrollbar()
        except:
            pass
    
    def _on_scrollbar_drag(self, event):
        """滚动条拖动"""
        self._on_scrollbar_click(event)
    
    def _btn_hover(self, is_hover):
        color = Theme.PRIMARY_HOVER if is_hover else Theme.PRIMARY
        self.run_btn.configure(bg=color)
        self.run_label.configure(bg=color)
    
    def _on_mode_change(self):
        mode = self.operation.get()
        enabled = mode in ('smart',)
        for card in self.preset_cards:
            card.set_enabled(enabled)
    
    def _open_custom_settings(self):
        """打开自定义设置窗口"""
        def on_save(settings):
            self.preset.set('custom')
            self.log_panel.log("自定义格式设置已保存", 'success')
        
        CustomSettingsDialog(self.root, on_save=on_save)
    
    def browse_input(self):
        is_windows = (os.name == 'nt')
        if is_windows:
            filetypes = [
                ("所有支持格式", "*.docx *.doc *.wps"),
                ("Word 文档 (.docx)", "*.docx"),
                ("Word 97-2003 (.doc)", "*.doc"),
                ("WPS 文档 (.wps)", "*.wps"),
                ("所有文件", "*.*"),
            ]
        else:
            filetypes = [
                ("Word 文档 (.docx)", "*.docx"),
                ("所有文件", "*.*"),
            ]
        filename = filedialog.askopenfilename(
            title="选择Word文档",
            filetypes=filetypes
        )
        if filename:
            self.input_file.set(filename)
            p = Path(filename)
            output_name = f"{p.stem}_processed{p.suffix}"
            self.output_file.set(str(p.parent / output_name))
            self.log_panel.log(f"已选择: {p.name}", 'info')
            self.log_panel.log(f"输出格式已自动设置为: {p.suffix or '.docx'}", 'info')
            self.result_panel.reset()
    
    def browse_output(self):
        is_windows = (os.name == 'nt')
        if is_windows:
            filetypes = [
                ("所有支持格式", "*.docx *.doc *.wps"),   # ← 默认选中，空格分隔
                ("Word 文档 (.docx)", "*.docx"),
                ("Word 97-2003 (.doc)", "*.doc"),
                ("WPS 文档 (.wps)", "*.wps"),
            ]
        else:
            filetypes = [
                ("Word 文档 (.docx)", "*.docx"),
            ]
        filename = filedialog.asksaveasfilename(
            title="保存为",
            defaultextension=".docx",
            filetypes=filetypes
        )
        if filename:
            self.output_file.set(filename)
    
    def run_operation(self):
        input_path = self.input_file.get().strip()
        output_path = self.output_file.get().strip()
        mode = self.operation.get()
        
        if not input_path:
            messagebox.showerror("提示", "请先选择输入文件")
            return
        
        if not os.path.exists(input_path):
            messagebox.showerror("错误", "文件不存在")
            return

        # Linux: 仅支持 .docx，.doc/.wps 需在 Windows 上转换
        if os.name != 'nt':
            input_ext = Path(input_path).suffix.lower()
            output_ext = Path(output_path).suffix.lower() if output_path else ''
            if input_ext in ('.doc', '.wps') or output_ext in ('.doc', '.wps'):
                messagebox.showerror(
                    "不支持的格式",
                    "Linux 版本仅支持 .docx 文件。.doc/.wps 请在 Windows 上转换，或先保存为 .docx 再处理。"
                )
                return
        
        if mode != 'analyze' and not output_path:
            messagebox.showerror("提示", "请指定输出文件")
            return
        
        self.run_btn.configure(bg=Theme.TEXT_MUTED)
        self.run_label.configure(bg=Theme.TEXT_MUTED, text="处理中...")
        
        thread = threading.Thread(
            target=self._do_operation,
            args=(input_path, output_path, mode)
        )
        thread.start()
    
    def _do_operation(self, input_path, output_path, mode):
        temp_docx = None
        temp_output_docx = None
        try:
            from docx import Document
            
            self.log_panel.log(f"\n{'─' * 35}", 'info')
            self.log_panel.log(f"开始处理: {Path(input_path).name}", 'info')
            
            ext = Path(input_path).suffix.lower()
            if ext in ('.doc', '.wps'):
                self.log_panel.log(f"检测到 {ext} 格式，正在转换...", 'info')
                from scripts.converter import convert_to_docx
                try:
                    temp_docx = convert_to_docx(input_path)
                except RuntimeError as e:
                    self.root.after(0, lambda: messagebox.showerror(
                        "转换失败",
                        "未检测到 WPS 或 Microsoft Office，请先安装后再试。"
                    ))
                    raise
                input_path = temp_docx
                self.log_panel.log("转换成功", 'success')
            
            output_ext = Path(output_path).suffix.lower()
            needs_convert_back = output_ext in ('.doc', '.wps')
            if needs_convert_back:
                import tempfile
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                    temp_output_docx = tmp.name
                output_path_docx = temp_output_docx
            else:
                output_path_docx = output_path
            
            if mode == 'analyze':
                doc = Document(input_path)
                results = {
                    'punctuation': analyze_punctuation(doc),
                    'numbering': analyze_numbering(doc),
                    'paragraph': analyze_paragraph_format(doc),
                    'font': analyze_font(doc)
                }
                self.root.after(0, lambda: self.result_panel.show_diagnosis(results))
                self.log_panel.log("诊断完成", 'success')
                
            elif mode == 'punctuation':
                self._run_punctuation(input_path, output_path_docx)
                self.root.after(0, lambda: self.result_panel.show_success(
                    "标点修复完成", Path(output_path).name
                ))
                
            elif mode == 'smart':
                import tempfile
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                    temp_path = tmp.name
                
                self.log_panel.log("步骤 1/2: 修复标点...", 'info')
                self._run_punctuation(input_path, temp_path, quiet=True)
                
                self.log_panel.log("步骤 2/2: 应用格式...", 'info')
                self._run_format(temp_path, output_path_docx)
                
                os.unlink(temp_path)
                
                self.root.after(0, lambda: self.result_panel.show_success(
                    "处理完成", Path(output_path).name
                ))
            
            if mode != 'analyze' and needs_convert_back:
                from scripts.converter import convert_from_docx
                try:
                    self.log_panel.log(f"正在转换回 {output_ext} 格式...", 'info')
                    actual_output = convert_from_docx(
                        output_path_docx, output_path,
                        format=output_ext.lstrip('.')
                    )
                    # convert_from_docx 可能回退到 .doc（当系统没有 WPS Office 时）
                    if actual_output and actual_output != output_path:
                        output_path = actual_output
                        self.log_panel.log(
                            f"保存 {output_ext} 需要安装 WPS Office，已自动保存为 .doc 格式",
                            'info'
                        )
                except RuntimeError as e:
                    if "未检测到" in str(e):
                        self.root.after(0, lambda: messagebox.showerror(
                            "转换失败",
                            "未检测到 WPS 或 Microsoft Office，请先安装后再试。"
                        ))
                        raise
                    # 其他 RuntimeError：回退保存为 .docx
                    self._fallback_to_docx(output_path, output_path_docx)
                    output_path = str(Path(output_path).with_suffix('.docx'))
                except Exception as e:
                    # COM 错误等：回退保存为 .docx
                    self.log_panel.log(f"转换回 {output_ext} 失败: {e}", 'info')
                    self._fallback_to_docx(output_path, output_path_docx)
                    output_path = str(Path(output_path).with_suffix('.docx'))
                finally:
                    if os.path.exists(output_path_docx) and output_path_docx != output_path:
                        try:
                            os.unlink(output_path_docx)
                        except Exception:
                            pass
            
            self.log_panel.log("全部完成", 'success')
            
            if mode != 'analyze':
                final_path = output_path  # 捕获到局部变量供 lambda 使用
                self.root.after(0, lambda: messagebox.showinfo(
                    "完成", f"文件已保存至:\n{final_path}"
                ))
        
        except Exception as e:
            error_msg = str(e)  # 先保存错误信息
            self.log_panel.log(f"错误: {error_msg}", 'error')
            import traceback
            self.log_panel.log(traceback.format_exc(), 'error')
            self.root.after(0, lambda msg=error_msg: messagebox.showerror("错误", msg))
        
        finally:
            if temp_docx and os.path.exists(temp_docx):
                os.unlink(temp_docx)
            if temp_output_docx and os.path.exists(temp_output_docx):
                os.unlink(temp_output_docx)
            self.root.after(0, self._reset_btn)
    
    def _fallback_to_docx(self, original_output_path, docx_source_path):
        """转换回原格式失败时，将已处理好的 .docx 直接保存"""
        import shutil
        fallback_path = str(Path(original_output_path).with_suffix('.docx'))
        try:
            shutil.copy2(docx_source_path, fallback_path)
            self.log_panel.log(
                f"已回退保存为 .docx 格式: {Path(fallback_path).name}", 'info'
            )
        except Exception as e:
            self.log_panel.log(f"回退保存也失败: {e}", 'error')
    
    def _reset_btn(self):
        self.run_btn.configure(bg=Theme.PRIMARY)
        self.run_label.configure(bg=Theme.PRIMARY, text="开始处理")
    
    def _run_punctuation(self, input_path, output_path, quiet=False):
        from docx import Document
        from scripts.punctuation import process_paragraph
        
        doc = Document(input_path)
        changes = 0
        
        for para in doc.paragraphs:
            if process_paragraph(para):
                changes += 1
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if process_paragraph(para):
                            changes += 1
        
        doc.save(output_path)
        if not quiet:
            self.log_panel.log(f"修复了 {changes} 处标点", 'success')
    
    def _run_format(self, input_path, output_path):
        preset_name = self.preset.get()
        
        import io
        old_stdout = sys.stdout
        sys.stdout = io.StringIO()
        
        try:
            format_document(input_path, output_path, preset_name)
        finally:
            sys.stdout = old_stdout
        
        if preset_name == 'custom':
            try:
                custom = load_custom_settings()
                preset_label = custom.get('name', '自定义格式') if custom else '自定义格式'
            except Exception:
                preset_label = '自定义格式'
        else:
            preset_label = PRESETS.get(preset_name, {}).get('name', preset_name)
        self.log_panel.log(f"应用格式: {preset_label}", 'success')


def main():
    try:
        from ctypes import windll
        windll.shcore.SetProcessDpiAwareness(1)
    except:
        pass
    
    root = tk.Tk()
    app = DocFormatApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
