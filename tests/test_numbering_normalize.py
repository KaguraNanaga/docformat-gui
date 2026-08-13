# -*- coding: utf-8 -*-
"""Automatic-numbering normalization regression tests."""

import zipfile

from docx import Document

from scripts.formatter import detect_para_type, format_document, normalize_automatic_numbering


_CONTENT_TYPES = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
<Override PartName="/word/numbering.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml"/>
</Types>'''

_RELS = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>'''

_DOC_RELS = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering" Target="numbering.xml"/>
</Relationships>'''

_NUMBERING = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:abstractNum w:abstractNumId="0">
<w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="chineseCounting"/><w:lvlText w:val="%1、"/></w:lvl>
<w:lvl w:ilvl="1"><w:start w:val="1"/><w:numFmt w:val="decimal"/><w:lvlText w:val="（%2）"/></w:lvl>
</w:abstractNum>
<w:num w:numId="1"><w:abstractNumId w:val="0"/></w:num>
</w:numbering>'''

_DOCUMENT = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body>
<w:p><w:pPr><w:numPr><w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr></w:pPr><w:r><w:t>一级标题</w:t></w:r></w:p>
<w:p><w:pPr><w:numPr><w:ilvl w:val="1"/><w:numId w:val="1"/></w:numPr></w:pPr><w:r><w:t>子项说明</w:t></w:r></w:p>
<w:p><w:pPr><w:numPr><w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr></w:pPr><w:r><w:t>二级内容</w:t></w:r></w:p>
<w:p><w:r><w:t>普通正文段落没有编号。</w:t></w:r></w:p>
</w:body></w:document>'''


def _write_numbered_docx(path, document_xml=_DOCUMENT, numbering_xml=_NUMBERING):
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("[Content_Types].xml", _CONTENT_TYPES)
        archive.writestr("_rels/.rels", _RELS)
        archive.writestr("word/document.xml", document_xml)
        archive.writestr("word/_rels/document.xml.rels", _DOC_RELS)
        archive.writestr("word/numbering.xml", numbering_xml)


def test_normalize_converts_auto_numbers_to_literal(tmp_path):
    source = tmp_path / "numbered.docx"
    _write_numbered_docx(source)
    document = Document(str(source))

    assert normalize_automatic_numbering(document) == 3
    texts = [paragraph.text for paragraph in document.paragraphs]
    assert texts[:3] == ["一、一级标题", "（1）子项说明", "二、二级内容"]
    assert all(
        paragraph._p.pPr is None or paragraph._p.pPr.numPr is None
        for paragraph in document.paragraphs
    )
    assert detect_para_type(texts[0], 0, 4, None, texts) == "heading1"
    assert detect_para_type(texts[1], 1, 4, None, texts) == "heading4"


def test_format_document_keeps_literal_numbers(tmp_path):
    source = tmp_path / "numbered.docx"
    target = tmp_path / "formatted.docx"
    _write_numbered_docx(source)
    format_document(str(source), str(target), preset_name="official")
    result = Document(str(target))
    texts = [paragraph.text for paragraph in result.paragraphs if paragraph.text.strip()]
    assert {"一、一级标题", "（1）子项说明", "二、二级内容"} <= set(texts)


def test_typed_number_and_numbering_property_do_not_double(tmp_path):
    document_xml = _DOCUMENT.replace(
        '<w:p><w:r><w:t>普通正文段落没有编号。</w:t></w:r></w:p>',
        '<w:p><w:pPr><w:numPr><w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr></w:pPr>'
        '<w:r><w:t>三、手打编号段落</w:t></w:r></w:p>',
    )
    source = tmp_path / "mixed.docx"
    _write_numbered_docx(source, document_xml)
    document = Document(str(source))
    normalize_automatic_numbering(document)
    assert document.paragraphs[3].text == "三、手打编号段落"


def test_automatic_bullet_is_preserved(tmp_path):
    numbering_xml = _NUMBERING.replace(
        '<w:num w:numId="1"><w:abstractNumId w:val="0"/></w:num>',
        '<w:abstractNum w:abstractNumId="2"><w:lvl w:ilvl="0">'
        '<w:start w:val="1"/><w:numFmt w:val="bullet"/><w:lvlText w:val="•"/>'
        '</w:lvl></w:abstractNum><w:num w:numId="1"><w:abstractNumId w:val="0"/></w:num>'
        '<w:num w:numId="2"><w:abstractNumId w:val="2"/></w:num>',
    )
    document_xml = _DOCUMENT.replace(
        '<w:p><w:r><w:t>普通正文段落没有编号。</w:t></w:r></w:p>',
        '<w:p><w:pPr><w:numPr><w:ilvl w:val="0"/><w:numId w:val="2"/></w:numPr>'
        '</w:pPr><w:r><w:t>项目符号内容</w:t></w:r></w:p>',
    )
    source = tmp_path / "bullets.docx"
    _write_numbered_docx(source, document_xml, numbering_xml)
    document = Document(str(source))
    assert normalize_automatic_numbering(document) == 3
    bullet = document.paragraphs[3]
    assert bullet.text == "项目符号内容"
    assert bullet._p.pPr.numPr is not None
