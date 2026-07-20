"""Chinese typography safeguards for Word output."""

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from scripts.east_asian_typography import apply_chinese_line_break_rules
from scripts.formatter import format_document
from scripts.punctuation import process_document


def _property_value(paragraph, local_name):
    p_pr = paragraph._p.pPr
    assert p_pr is not None
    element = p_pr.find(qn(f"w:{local_name}"))
    assert element is not None
    return element.get(qn("w:val"))


def test_applies_kinsoku_to_body_nested_tables_headers_and_footers():
    document = Document()
    body = document.add_paragraph("正文，必须避免标点起行。")
    cell = document.add_table(rows=1, cols=1).cell(0, 0)
    cell_paragraph = cell.paragraphs[0]
    cell_paragraph.text = "表格内的正文，也应遵守中文换行规则。"
    nested = cell.add_table(rows=1, cols=1).cell(0, 0).paragraphs[0]
    nested.text = "嵌套表格内也不能让句号起行。"
    header = document.sections[0].header.paragraphs[0]
    header.text = "页眉，中文排版。"
    footer = document.sections[0].footer.paragraphs[0]
    footer.text = "页脚，中文排版。"

    assert apply_chinese_line_break_rules(document) == 5
    for paragraph in (body, cell_paragraph, nested, header, footer):
        assert _property_value(paragraph, "kinsoku") == "1"
        assert _property_value(paragraph, "overflowPunct") == "0"

    # Reapplying must not add a second set of OOXML properties.
    assert apply_chinese_line_break_rules(document) == 0


def test_existing_overflow_punctuation_setting_is_disabled():
    document = Document()
    paragraph = document.add_paragraph("已有模板设置也要以中文禁则为准。")
    overflow = OxmlElement("w:overflowPunct")
    overflow.set(qn("w:val"), "1")
    paragraph._p.get_or_add_pPr().append(overflow)

    assert apply_chinese_line_break_rules(document) == 1
    assert _property_value(paragraph, "kinsoku") == "1"
    assert _property_value(paragraph, "overflowPunct") == "0"


def test_formatter_and_punctuation_only_write_line_break_rules(tmp_path):
    source = tmp_path / "source.docx"
    formatted_path = tmp_path / "formatted.docx"
    punctuated_path = tmp_path / "punctuated.docx"
    document = Document()
    document.add_paragraph("关于开展重点工作的通知")
    document.add_paragraph("请按中文习惯控制标点,避免标点起行.")
    document.save(source)

    format_document(
        str(source), str(formatted_path), preset_name="official",
        custom_settings={"page_number": False},
    )
    formatted = Document(formatted_path)
    paragraph = next(p for p in formatted.paragraphs if "避免标点起行" in p.text)
    assert _property_value(paragraph, "kinsoku") == "1"

    process_document(str(source), str(punctuated_path))
    fixed = Document(punctuated_path).paragraphs[1]
    assert fixed.text == "请按中文习惯控制标点，避免标点起行。"
    assert _property_value(fixed, "overflowPunct") == "0"
