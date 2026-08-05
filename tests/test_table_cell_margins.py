"""Custom table-cell margin regression tests."""

from copy import deepcopy

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm

import docformat_gui
from scripts.formatter import PRESETS, format_document


def test_table_cell_margin_defaults_and_custom_values_round_trip(tmp_path, monkeypatch):
    defaults = docformat_gui.DEFAULT_CUSTOM_SETTINGS["table"]
    assert defaults["cell_margin_top_cm"] == 0.0
    assert defaults["cell_margin_bottom_cm"] == 0.0
    assert defaults["cell_margin_left_cm"] == 0.05
    assert defaults["cell_margin_right_cm"] == 0.05

    preset = deepcopy(docformat_gui.DEFAULT_CUSTOM_SETTINGS)
    preset.update({"id": "table-margins", "name": "表格边距"})
    preset["table"].update({
        "cell_margin_top_cm": 0.1,
        "cell_margin_bottom_cm": 0.2,
        "cell_margin_left_cm": 0.3,
        "cell_margin_right_cm": 0.4,
    })
    config_file = tmp_path / "custom_settings.json"
    monkeypatch.setattr(docformat_gui, "CONFIG_FILE", config_file)
    docformat_gui.save_custom_settings({
        "schema_version": docformat_gui.CONFIG_SCHEMA_VERSION,
        "active_preset_id": "table-margins",
        "presets": [preset],
    })

    loaded = docformat_gui.get_active_user_preset(docformat_gui.load_custom_settings())
    assert loaded["table"]["cell_margin_top_cm"] == 0.1
    assert loaded["table"]["cell_margin_bottom_cm"] == 0.2
    assert loaded["table"]["cell_margin_left_cm"] == 0.3
    assert loaded["table"]["cell_margin_right_cm"] == 0.4


def test_custom_table_cell_margins_are_written_to_docx(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "内容"
    document.save(source)

    settings = deepcopy(PRESETS["official"])
    settings["page_number"] = False
    settings.setdefault("table", {}).update({
        "cell_margin_top_cm": 0.1,
        "cell_margin_bottom_cm": 0.2,
        "cell_margin_left_cm": 0.3,
        "cell_margin_right_cm": 0.4,
    })
    format_document(
        str(source), str(output), preset_name="custom", custom_settings=settings,
    )

    result_table = Document(output).tables[0]
    margins = result_table._tbl.tblPr.find(qn("w:tblCellMar"))
    assert margins is not None
    for side, value in {
        "top": 0.1,
        "bottom": 0.2,
        "left": 0.3,
        "right": 0.4,
    }.items():
        assert int(margins.find(qn(f"w:{side}")).get(qn("w:w"))) == int(Cm(value).twips)
