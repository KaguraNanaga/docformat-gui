"""Regression tests for media preservation and standalone attachment markers."""

import struct
import zlib

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt

from scripts.formatter import format_document


def _write_tiny_png(path):
    def chunk(kind, data):
        return struct.pack(">I", len(data)) + kind + data + struct.pack(
            ">I", zlib.crc32(kind + data) & 0xFFFFFFFF
        )

    raw = b"\x00\xcc\x33\x33"
    path.write_bytes(
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0))
        + chunk(b"IDAT", zlib.compress(raw))
        + chunk(b"IEND", b"")
    )


def _has_drawing(paragraph):
    return bool(paragraph._p.xpath(".//w:drawing"))


def test_format_document_preserves_mixed_and_standalone_media(tmp_path):
    image_path = tmp_path / "tiny.png"
    _write_tiny_png(image_path)
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("关于图片保留的报告")
    mixed = doc.add_paragraph()
    mixed.add_run("一是")
    mixed.add_run().add_picture(str(image_path), width=Inches(0.2))
    mixed.add_run("推进重点工作。")
    media_only = doc.add_paragraph()
    media_only.add_run().add_picture(str(image_path), width=Inches(0.2))
    media_only.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    media_only.paragraph_format.line_spacing = Pt(28)
    # Legacy VML and OLE/object nodes must also survive a run-rebuilding path.
    mixed.add_run()._r.append(OxmlElement("w:pict"))
    mixed.add_run()._r.append(OxmlElement("w:object"))
    doc.save(source)

    format_document(
        str(source), str(output), preset_name="official",
        custom_settings={"page_number": False, "bold_serial": True},
    )
    formatted = Document(output)
    assert len(formatted.inline_shapes) == 2
    media_paragraphs = [p for p in formatted.paragraphs if _has_drawing(p)]
    assert len(media_paragraphs) == 2
    assert all(p.paragraph_format.line_spacing_rule == WD_LINE_SPACING.SINGLE for p in media_paragraphs)
    xml = "".join(p._p.xml for p in formatted.paragraphs)
    assert "<w:pict" in xml
    assert "<w:object" in xml


def test_standalone_attachment_starts_new_page_and_formats_following_title(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("关于附件格式的通知")
    doc.add_paragraph("正文内容。")
    doc.add_paragraph("附件1")
    doc.add_paragraph("专项工作实施方案")
    doc.add_paragraph("一、工作目标")
    doc.save(source)

    format_document(str(source), str(output), preset_name="official", custom_settings={"page_number": False})
    formatted = Document(output)
    marker = next(p for p in formatted.paragraphs if p.text.strip() == "附件1")
    title = next(p for p in formatted.paragraphs if p.text.strip() == "专项工作实施方案")
    heading = next(p for p in formatted.paragraphs if p.text.strip() == "一、工作目标")
    assert marker.paragraph_format.page_break_before is True
    assert marker.alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert marker.paragraph_format.left_indent == Pt(0)
    assert marker.paragraph_format.first_line_indent == Pt(0)
    assert title.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert heading.alignment == WD_ALIGN_PARAGRAPH.LEFT
