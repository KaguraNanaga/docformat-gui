"""Table lead-in and cell-option regression tests."""

from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

from scripts import formatter


def _settings(**table_overrides):
    settings = deepcopy(formatter.PRESETS["official"])
    settings["page_number"] = False
    settings["table"] = {
        "font_cn": "仿宋_GB2312", "font_en": "Times New Roman", "size": 12,
        "line_spacing": 22, "header_bold": True, "optimize": True,
        "before_table_blank_line": False, "after_table_blank_line": False,
        **table_overrides,
    }
    return settings


def _source_table(path):
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "序号"
    table.cell(0, 1).text = "内容"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "需要适应单元格的文字"
    document.save(path)


def test_cell_margins_and_fit_text_are_written_per_cell(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    _source_table(source)
    margins = {"top": 0.03, "bottom": 0.04, "left": 0.05, "right": 0.06}
    formatter.format_document(
        str(source), str(output), preset_name="custom",
        custom_settings=_settings(
            cell_margins_same_as_table=False,
            cell_margin_top_cm=margins["top"],
            cell_margin_bottom_cm=margins["bottom"],
            cell_margin_left_cm=margins["left"],
            cell_margin_right_cm=margins["right"],
            cell_fit_text=True,
        ),
    )
    for row in Document(output).tables[0].rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.find(qn("w:tcMar"))
            assert tc_mar is not None
            for side, value in margins.items():
                assert int(tc_mar.find(qn(f"w:{side}")).get(qn("w:w"))) == int(Cm(value).twips)
            assert tc_pr.find(qn("w:tcFitText")) is not None


def test_fit_text_can_be_removed_from_existing_cells(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    _source_table(source)
    document = Document(source)
    for row in document.tables[0].rows:
        for cell in row.cells:
            cell._tc.get_or_add_tcPr().append(OxmlElement("w:tcFitText"))
    document.save(source)
    formatter.format_document(
        str(source), str(output), preset_name="custom",
        custom_settings=_settings(cell_fit_text=False),
    )
    for row in Document(output).tables[0].rows:
        for cell in row.cells:
            assert cell._tc.get_or_add_tcPr().find(qn("w:tcFitText")) is None


def test_cell_options_apply_when_generic_optimization_is_off(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    _source_table(source)
    settings = _settings(
        cell_margin_top_cm=0.1, cell_margin_bottom_cm=0.1,
        cell_margin_left_cm=0.2, cell_margin_right_cm=0.2,
    )
    settings["layout_mode"] = "generic"
    settings["table"]["optimize"] = False
    formatter.format_document(str(source), str(output), "custom", custom_settings=settings)
    table = Document(output).tables[0]
    table_margins = table._tbl.tblPr.find(qn("w:tblCellMar"))
    assert int(table_margins.find(qn("w:left")).get(qn("w:w"))) == int(Cm(0.2).twips)


def test_table_lead_in_is_body_and_alignment_is_repaired(tmp_path, monkeypatch):
    lead_in_text = "现申请有关事项变更，具体情况如下："
    texts = ["情况说明", "有关单位：", lead_in_text, "某某公司", "2026年2月5日"]
    assert formatter.detect_para_type(
        lead_in_text, 10, 18, WD_ALIGN_PARAGRAPH.RIGHT, texts,
        all_texts_index=2, prev_para_type="body",
    ) == "body"

    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("关于变更情况的说明")
    document.add_paragraph("有关单位：")
    lead_in = document.add_paragraph(lead_in_text)
    lead_in.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph("")
    document.add_table(rows=1, cols=2)
    document.add_paragraph("某某公司")
    document.add_paragraph("2026年2月5日")
    document.save(source)

    original_detect = formatter.detect_para_type

    def misclassify(text, *args, **kwargs):
        return "signature" if text.strip() == lead_in_text else original_detect(text, *args, **kwargs)

    monkeypatch.setattr(formatter, "detect_para_type", misclassify)
    formatter.format_document(str(source), str(output), preset_name="official")
    paragraph = next(item for item in Document(output).paragraphs if item.text == lead_in_text)
    assert paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert paragraph.paragraph_format.first_line_indent.pt > 0
