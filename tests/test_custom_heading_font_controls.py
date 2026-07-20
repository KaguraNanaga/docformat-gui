"""Heading Chinese and English/numeric font controls."""

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from scripts.formatter import PRESETS, format_document


PROJECT_ROOT = Path(__file__).parent.parent


def test_custom_settings_expose_each_heading_font_choice_in_the_ui():
    source = (PROJECT_ROOT / "docformat_gui.py").read_text(encoding="utf-8")
    for level in range(1, 5):
        assert f"self.h{level}_font_var = tk.StringVar()" in source
        assert f"self.h{level}_font_en_var = tk.StringVar()" in source


def test_each_heading_writes_chinese_and_english_fonts_to_ooxml(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("关于标题字体的通知")
    doc.add_paragraph("一、Heading One")
    doc.add_paragraph("（二）Heading Two")
    doc.add_paragraph("3. Heading Three")
    doc.add_paragraph("（4）Heading Four")
    doc.save(source)

    custom = deepcopy(PRESETS["official"])
    for level in range(1, 5):
        custom[f"heading{level}"] = {
            **custom[f"heading{level}"],
            "font_cn": f"Chinese{level}",
            "font_en": f"Latin{level}",
        }
    custom["page_number"] = False
    format_document(str(source), str(output), preset_name="custom", custom_settings=custom)

    formatted = Document(output)
    headings = (
        "一、Heading One", "（二）Heading Two", "3. Heading Three", "（4）Heading Four",
    )
    for level, heading in enumerate(headings, start=1):
        paragraph = next(p for p in formatted.paragraphs if p.text.strip() == heading)
        fonts = paragraph.runs[0]._r.rPr.find(qn("w:rFonts"))
        assert fonts.get(qn("w:eastAsia")) == f"Chinese{level}"
        assert fonts.get(qn("w:ascii")) == f"Latin{level}"
        assert fonts.get(qn("w:hAnsi")) == f"Latin{level}"
