"""General-purpose layout and local sample-learning regression tests."""

from copy import deepcopy
import json

import pytest
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

from scripts.formatter import PRESETS, format_document
from scripts.layout_analyzer import LayoutAnalysisError, analyze_reference_layout


def _set_run_font(run, name, size, *, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    fonts = run._r.get_or_add_rPr().find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        run._r.get_or_add_rPr().insert(0, fonts)
    fonts.set(qn("w:eastAsia"), name)
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")


def _generic_settings():
    settings = deepcopy(PRESETS["official"])
    settings.update({
        "name": "通用材料", "layout_mode": "generic", "remove_background": False,
        "page_number": False, "deep_clean": False, "split_heading_at_punct": False,
        "first_line_bold": False, "bold_serial": False,
    })
    settings["page"] = {
        "top": 2.2, "bottom": 2.2, "left": 2.4, "right": 2.4,
        "width_mm": 210, "height_mm": 297, "orientation": "landscape",
    }
    settings["title"] = {
        "font_cn": "黑体", "font_en": "Arial", "size": 20, "bold": True,
        "align": "center", "indent": 0, "line_spacing": 30,
        "space_before": 0, "space_after": 8,
    }
    settings["heading1"] = {
        "font_cn": "黑体", "font_en": "Arial", "size": 16, "bold": True,
        "align": "left", "indent": 0, "line_spacing": 26,
        "space_before": 8, "space_after": 4,
    }
    settings["body"] = {
        "font_cn": "宋体", "font_en": "Arial", "size": 14, "bold": False,
        "align": "justify", "indent": 28, "line_spacing": 24,
        "space_before": 0, "space_after": 0,
    }
    for role in ("recipient", "signature", "date", "attachment", "closing", "source_note"):
        settings[role] = deepcopy(settings["body"])
    settings["table"] = {
        "font_cn": "宋体", "font_en": "Arial", "size": 11, "bold": False,
        "line_spacing": 18, "first_line_indent": 0, "header_bold": False,
        "smart_align": False, "optimize": False,
        "before_table_blank_line": False, "after_table_blank_line": False,
    }
    return settings


def test_generic_layout_preserves_structure_inline_emphasis_and_orientation(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("学习材料", style="Title")
    document.add_paragraph("一、学习主题", style="Heading 1")
    body = document.add_paragraph()
    body.add_run("本次学习的")
    emphasized = body.add_run("核心指标")
    emphasized.bold = True
    emphasized.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    marked = body.add_run("需要重点研讨。")
    marked.font.highlight_color = WD_COLOR_INDEX.YELLOW
    document.add_paragraph("交流研讨：")
    document.add_paragraph("二○二六年七月二十八日")
    blank = document.add_paragraph("")
    blank.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    blank.paragraph_format.line_spacing = Pt(18)
    document.add_table(rows=1, cols=2)
    document.add_paragraph("研讨结束后形成纪要。")
    document.save(source)

    before_count = len(Document(source).paragraphs)
    format_document(str(source), str(output), "custom", custom_settings=_generic_settings())
    result = Document(output)
    assert len(result.paragraphs) == before_count
    assert result.sections[0].page_width > result.sections[0].page_height
    assert result.sections[0].orientation == WD_ORIENT.LANDSCAPE
    assert sum(not paragraph.text.strip() for paragraph in result.paragraphs) == 1
    colon = next(paragraph for paragraph in result.paragraphs if paragraph.text == "交流研讨：")
    date = next(paragraph for paragraph in result.paragraphs if paragraph.text.startswith("二○二六年"))
    assert colon.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert date.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    formatted_body = next(paragraph for paragraph in result.paragraphs if "核心指标" in paragraph.text)
    bold_run = next(run for run in formatted_body.runs if "核心指标" in run.text)
    highlighted = next(run for run in formatted_body.runs if "需要重点" in run.text)
    assert bold_run.bold is True
    assert bold_run.font.color.rgb == RGBColor(0xC0, 0x00, 0x00)
    assert highlighted.font.highlight_color == WD_COLOR_INDEX.YELLOW


def test_reference_learning_extracts_layout_without_copying_content(tmp_path):
    reference = tmp_path / "reference.docx"
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(2.1)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(title.add_run("不应复制的样例标题"), "黑体", 20, bold=True)
    heading = document.add_paragraph("一、学习安排", style="Heading 1")
    _set_run_font(heading.runs[0], "黑体", 16, bold=True)
    for text in ("不应进入预设的正文内容。", "另一段用于判断正文格式。"):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Pt(30)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(26)
        _set_run_font(paragraph.add_run(text), "仿宋", 15)
    document.save(reference)

    analysis = analyze_reference_layout(reference, base_settings=_generic_settings())
    settings = analysis.settings
    assert settings["layout_mode"] == "generic"
    assert settings["page"]["top"] == pytest.approx(1.8, abs=0.02)
    assert settings["title"]["size"] == pytest.approx(20)
    assert settings["body"]["indent"] == pytest.approx(30)
    serialized = json.dumps(settings, ensure_ascii=False)
    assert "不应复制" not in serialized
    assert "不应进入预设" not in serialized


def test_reference_learning_rejects_non_docx(tmp_path):
    path = tmp_path / "reference.txt"
    path.write_text("not a word document", encoding="utf-8")
    with pytest.raises(LayoutAnalysisError, match="docx"):
        analyze_reference_layout(path)


def test_generic_mode_strips_leading_spaces_only_when_style_indents(tmp_path):
    source = tmp_path / "spaces.docx"
    output = tmp_path / "spaces-out.docx"
    document = Document()
    document.add_paragraph("学习材料", style="Title")
    document.add_paragraph("　　各单位要结合实际认真落实。")
    document.save(source)
    format_document(str(source), str(output), "custom", custom_settings=_generic_settings())
    body = next(paragraph for paragraph in Document(output).paragraphs if "各单位" in paragraph.text)
    assert body.text.startswith("各单位")
    assert body._p.pPr.find(qn("w:ind")).get(qn("w:firstLineChars")) == "200"
